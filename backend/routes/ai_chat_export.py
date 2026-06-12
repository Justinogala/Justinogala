"""
AI Chat — Export conversations to Markdown, PDF, and DOCX.
"""
import re
import io
from datetime import datetime
from fastapi import Query, HTTPException, Depends
from fastapi.responses import Response

from config import db
from routes.auth import get_current_user


def _safe_filename(title: str) -> str:
    safe = re.sub(r'[^\w\s-]', '', title)[:60].strip()
    return safe or "chat-export"


def _build_markdown(title: str, created: str, msgs: list) -> str:
    lines = [f"# {title}\n"]
    if created:
        try:
            dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
            lines.append(f"*Exported from Munal AI &mdash; {dt.strftime('%B %d, %Y')}*\n")
        except Exception:
            lines.append(f"*Exported from Munal AI*\n")
    lines.append("---\n")

    for msg in msgs:
        role = "You" if msg["role"] == "user" else "Munal AI"
        lines.append(f"### {role}\n")
        lines.append(f"{msg.get('content', '')}\n")
        lines.append("")

    lines.append("---\n*Exported from Munal AI*")
    return "\n".join(lines)


def _build_pdf(title: str, created: str, msgs: list) -> bytes:
    import fitz  # PyMuPDF

    doc = fitz.open()
    WIDTH, HEIGHT = 595, 842  # A4
    MARGIN = 50
    usable_w = WIDTH - 2 * MARGIN
    y = MARGIN

    def new_page():
        nonlocal y
        page = doc.new_page(width=WIDTH, height=HEIGHT)
        y = MARGIN
        return page

    page = new_page()

    # Title
    y += 10
    page.insert_text((MARGIN, y), title[:80], fontsize=18, fontname="helv", color=(0.29, 0.27, 0.53))
    y += 28

    # Date
    date_str = "Exported from Munal AI"
    if created:
        try:
            dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
            date_str = f"Exported from Munal AI — {dt.strftime('%B %d, %Y')}"
        except Exception:
            pass
    page.insert_text((MARGIN, y), date_str, fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))
    y += 20

    page.draw_line((MARGIN, y), (WIDTH - MARGIN, y), color=(0.85, 0.85, 0.85), width=0.5)
    y += 15

    for msg in msgs:
        role = "You" if msg["role"] == "user" else "Munal AI"
        content = msg.get("content", "")
        role_color = (0.29, 0.27, 0.53) if msg["role"] == "assistant" else (0.2, 0.2, 0.2)

        if y > HEIGHT - 80:
            page = new_page()
        page.insert_text((MARGIN, y), role, fontsize=11, fontname="helvetica-bold", color=role_color)
        y += 18

        lines = []
        for paragraph in content.split('\n'):
            if not paragraph.strip():
                lines.append("")
                continue
            words = paragraph.split(' ')
            current_line = ""
            for word in words:
                test_line = f"{current_line} {word}".strip() if current_line else word
                text_width = fitz.get_text_length(test_line, fontname="helv", fontsize=10)
                if text_width > usable_w:
                    lines.append(current_line)
                    current_line = word
                else:
                    current_line = test_line
            if current_line:
                lines.append(current_line)

        for line in lines:
            if y > HEIGHT - MARGIN:
                page = new_page()
            page.insert_text((MARGIN, y), line, fontsize=10, fontname="helv", color=(0.15, 0.15, 0.15))
            y += 14

        y += 12

    if y > HEIGHT - 40:
        page = new_page()
    page.draw_line((MARGIN, HEIGHT - 40), (WIDTH - MARGIN, HEIGHT - 40), color=(0.85, 0.85, 0.85), width=0.5)
    page.insert_text((MARGIN, HEIGHT - 28), "Exported from Munal AI", fontsize=8, fontname="helv", color=(0.6, 0.6, 0.6))

    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _build_docx(title: str, created: str, msgs: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(75, 69, 135)

    date_str = "Exported from Munal AI"
    if created:
        try:
            dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
            date_str = f"Exported from Munal AI — {dt.strftime('%B %d, %Y')}"
        except Exception:
            pass
    date_para = doc.add_paragraph(date_str)
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("").paragraph_format.space_after = Pt(4)

    for msg in msgs:
        role = "You" if msg["role"] == "user" else "Munal AI"
        content = msg.get("content", "")

        role_para = doc.add_paragraph()
        role_run = role_para.add_run(role)
        role_run.bold = True
        role_run.font.size = Pt(11)
        if msg["role"] == "assistant":
            role_run.font.color.rgb = RGBColor(75, 69, 135)
        role_para.paragraph_format.space_after = Pt(2)

        for paragraph_text in content.split('\n'):
            p = doc.add_paragraph(paragraph_text)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)

        doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    footer_para = doc.add_paragraph("Exported from Munal AI")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(160, 160, 160)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def export_conversation_handler(conv_id: str, format: str, user: dict):
    """Export a conversation to MD, PDF, or DOCX."""
    conv = await db.ai_conversations.find_one({"id": conv_id, "user_id": user["id"]}, {"_id": 0})
    if not conv:
        raise HTTPException(404, "Conversation not found")

    msgs = await db.ai_messages.find(
        {"conversation_id": conv_id, "role": {"$in": ["user", "assistant"]}},
        {"_id": 0}
    ).sort("created_at", 1).to_list(length=500)

    title = conv.get("title", "Chat Export")
    created = conv.get("created_at", "")

    if format == "md":
        md = _build_markdown(title, created, msgs)
        return Response(
            content=md.encode("utf-8"),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{_safe_filename(title)}.md"'}
        )
    elif format == "pdf":
        pdf_bytes = _build_pdf(title, created, msgs)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{_safe_filename(title)}.pdf"'}
        )
    elif format == "docx":
        docx_bytes = _build_docx(title, created, msgs)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{_safe_filename(title)}.docx"'}
        )
