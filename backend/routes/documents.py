"""
Documents routes - CRUD for rich text documents.
"""
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from datetime import datetime, timezone
from typing import Optional
import uuid
import jwt
import os

from config import db, JWT_SECRET_KEY, JWT_ALGORITHM, logger

router = APIRouter(prefix="/documents", tags=["Documents"])
security = HTTPBearer(auto_error=False)


def _get_user_id(creds: HTTPAuthorizationCredentials):
    if not creds:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
        return payload.get("user_id") or payload.get("sub")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")


@router.get("")
async def list_documents(
    search: Optional[str] = Query(None),
    workspace_id: Optional[str] = Query(None),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """List all documents for the current user, optionally filtered by workspace"""
    user_id = _get_user_id(credentials)
    if workspace_id:
        query = {"$or": [
            {"user_id": user_id, "workspace_id": workspace_id, "deleted": {"$ne": True}},
            {"linked_workspaces": workspace_id, "deleted": {"$ne": True}},
        ]}
    else:
        query = {"user_id": user_id, "deleted": {"$ne": True}, "workspace_id": {"$in": [None, ""]}}
    if search:
        query["title"] = {"$regex": search, "$options": "i"}

    docs = await db.documents.find(query, {"_id": 0}).sort("updated_at", -1).to_list(200)
    return docs


@router.post("")
async def create_document(
    body: dict,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Create a new document"""
    user_id = _get_user_id(credentials)
    doc_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    doc = {
        "id": doc_id,
        "user_id": user_id,
        "workspace_id": body.get("workspace_id") or None,
        "title": body.get("title", "Untitled Document"),
        "content": body.get("content", ""),
        "template": body.get("template"),
        "word_count": len((body.get("content", "")).split()),
        "created_at": now,
        "updated_at": now,
        "deleted": False,
    }
    await db.documents.insert_one(doc)
    doc.pop("_id", None)
    return doc


@router.get("/{doc_id}")
async def get_document(
    doc_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Get a single document"""
    user_id = _get_user_id(credentials)
    doc = await db.documents.find_one({"id": doc_id, "user_id": user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.put("/{doc_id}")
async def update_document(
    doc_id: str,
    body: dict,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Update a document"""
    user_id = _get_user_id(credentials)
    existing = await db.documents.find_one({"id": doc_id, "user_id": user_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Document not found")

    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if "title" in body:
        update_data["title"] = body["title"]
    if "content" in body:
        update_data["content"] = body["content"]
        update_data["word_count"] = len(body["content"].split())

    await db.documents.update_one({"id": doc_id}, {"$set": update_data})
    updated = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    return updated


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Soft-delete a document"""
    user_id = _get_user_id(credentials)
    result = await db.documents.update_one(
        {"id": doc_id, "user_id": user_id},
        {"$set": {"deleted": True, "deleted_at": datetime.now(timezone.utc).isoformat(), "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"success": True}


@router.post("/{doc_id}/duplicate")
async def duplicate_document(
    doc_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Duplicate a document"""
    user_id = _get_user_id(credentials)
    original = await db.documents.find_one({"id": doc_id, "user_id": user_id}, {"_id": 0})
    if not original:
        raise HTTPException(status_code=404, detail="Document not found")

    new_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    new_doc = {
        **original,
        "id": new_id,
        "title": f"{original['title']} (Copy)",
        "created_at": now,
        "updated_at": now,
    }
    await db.documents.insert_one(new_doc)
    new_doc.pop("_id", None)
    return new_doc


@router.post("/ai-generate")
async def ai_generate_document(
    body: dict,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Generate a document using AI from a text prompt"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage

    user_id = _get_user_id(credentials)
    prompt = body.get("prompt", "").strip()
    if not prompt:
        raise HTTPException(status_code=400, detail="Prompt is required")

    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        raise HTTPException(status_code=503, detail="AI service not configured")

    system_msg = """You are a professional document writer. Generate well-structured HTML content for a document based on the user's request. 
Use proper HTML tags: <h1> for title, <h2> for sections, <h3> for subsections, <p> for paragraphs, <ul>/<ol>/<li> for lists, <strong> for bold, <em> for italic, <table>/<tr>/<th>/<td> for tables, <blockquote> for quotes.
Make the content comprehensive, professional, and ready to use. Do NOT include <html>, <head>, or <body> tags - just the content HTML."""

    try:
        chat = LlmChat(
            api_key=api_key,
            session_id=f"doc-gen-{uuid.uuid4()}",
            system_message=system_msg
        ).with_model("openai", "gpt-5.2")

        user_message = UserMessage(text=f"Create a professional document: {prompt}")
        content = await chat.send_message(user_message)

        # Extract title from first heading or prompt
        import re
        title_match = re.search(r'<h1[^>]*>(.*?)</h1>', content)
        title = title_match.group(1) if title_match else prompt[:60]

        # Save as new document
        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        doc = {
            "id": doc_id,
            "user_id": user_id,
            "title": title,
            "content": content,
            "template": "ai-generated",
            "word_count": len(content.split()),
            "created_at": now,
            "updated_at": now,
            "deleted": False,
        }
        await db.documents.insert_one(doc)
        doc.pop("_id", None)
        return doc

    except Exception as e:
        logger.error(f"AI document generation failed: {e}")
        # Fallback: create a basic document with the prompt
        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        fallback_content = f"<h1>{prompt}</h1><p>Start writing your document here...</p><h2>Overview</h2><p></p><h2>Details</h2><p></p><h2>Conclusion</h2><p></p>"
        doc = {
            "id": doc_id,
            "user_id": user_id,
            "title": prompt[:60],
            "content": fallback_content,
            "template": "ai-generated",
            "word_count": len(fallback_content.split()),
            "created_at": now,
            "updated_at": now,
            "deleted": False,
        }
        await db.documents.insert_one(doc)
        doc.pop("_id", None)
        return doc


@router.get("/{doc_id}/export/docx")
async def export_docx(
    doc_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Export a document as DOCX"""
    from fastapi.responses import Response
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import re
    from html.parser import HTMLParser

    user_id = _get_user_id(credentials)
    doc = await db.documents.find_one({"id": doc_id, "user_id": user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Parse HTML to DOCX
    docx_doc = DocxDocument()
    html_content = doc.get("content", "")

    class HTMLToDocx(HTMLParser):
        def __init__(self, document):
            super().__init__()
            self.doc = document
            self.current_para = None
            self.current_run = None
            self.bold = False
            self.italic = False
            self.underline = False
            self.strike = False
            self.in_heading = 0
            self.in_list = False
            self.in_ordered = False
            self.list_counter = 0
            self.in_blockquote = False
            self.in_table = False
            self.table = None
            self.current_row = None
            self.current_cell = None
            self.text_align = None

        def handle_starttag(self, tag, attrs):
            attrs_dict = dict(attrs)
            style = attrs_dict.get('style', '')

            if tag in ('h1', 'h2', 'h3'):
                level = int(tag[1])
                self.in_heading = level
                self.current_para = self.doc.add_heading('', level=level)
            elif tag == 'p':
                if self.in_table and self.current_cell:
                    self.current_para = self.current_cell.paragraphs[0] if self.current_cell.paragraphs else self.current_cell.add_paragraph()
                else:
                    self.current_para = self.doc.add_paragraph()
                if 'text-align: center' in style or 'text-align:center' in style:
                    self.current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in style or 'text-align:right' in style:
                    self.current_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'text-align: justify' in style or 'text-align:justify' in style:
                    self.current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif tag == 'strong' or tag == 'b':
                self.bold = True
            elif tag == 'em' or tag == 'i':
                self.italic = True
            elif tag == 'u':
                self.underline = True
            elif tag == 's' or tag == 'del':
                self.strike = True
            elif tag == 'ul':
                self.in_list = True
                self.in_ordered = False
            elif tag == 'ol':
                self.in_list = True
                self.in_ordered = True
                self.list_counter = 0
            elif tag == 'li':
                if self.in_ordered:
                    self.list_counter += 1
                    self.current_para = self.doc.add_paragraph(f'{self.list_counter}. ', style='List Number')
                else:
                    self.current_para = self.doc.add_paragraph('', style='List Bullet')
            elif tag == 'blockquote':
                self.in_blockquote = True
            elif tag == 'table':
                self.in_table = True
                self.table = self.doc.add_table(rows=0, cols=0)
                self.table.style = 'Table Grid'
            elif tag == 'tr':
                if self.table:
                    self.current_row = self.table.add_row() if self.table.rows else None
                    if not self.current_row:
                        self.table.add_row()
                        self.current_row = self.table.rows[-1]
            elif tag in ('td', 'th'):
                if self.current_row:
                    if len(self.current_row.cells) == 0 or self.current_cell is None:
                        # Need to add column
                        if len(self.table.columns) < len(self.current_row.cells) + 1:
                            self.table.add_column(Inches(2))
                        idx = 0 if self.current_cell is None else len([c for c in self.current_row.cells]) - 1
                    self.current_cell = self.current_row.cells[min(len(self.current_row.cells)-1, 0)] if self.current_row.cells else None
                    self.current_para = None
            elif tag == 'br':
                if self.current_para:
                    self.current_para.add_run('\n')
            elif tag == 'hr':
                para = self.doc.add_paragraph()
                para.add_run('─' * 50)
            elif tag == 'img':
                src = attrs_dict.get('src', '')
                if src:
                    para = self.doc.add_paragraph()
                    para.add_run(f'[Image: {src}]')

        def handle_endtag(self, tag):
            if tag in ('h1', 'h2', 'h3'):
                self.in_heading = 0
                self.current_para = None
            elif tag == 'p':
                self.current_para = None
            elif tag == 'strong' or tag == 'b':
                self.bold = False
            elif tag == 'em' or tag == 'i':
                self.italic = False
            elif tag == 'u':
                self.underline = False
            elif tag == 's' or tag == 'del':
                self.strike = False
            elif tag in ('ul', 'ol'):
                self.in_list = False
                self.in_ordered = False
            elif tag == 'blockquote':
                self.in_blockquote = False
            elif tag == 'table':
                self.in_table = False
                self.table = None
                self.current_row = None
                self.current_cell = None
            elif tag == 'tr':
                self.current_row = None
                self.current_cell = None
            elif tag in ('td', 'th'):
                self.current_cell = None

        def handle_data(self, data):
            text = data.strip()
            if not text:
                return
            if self.current_para is None:
                if self.in_blockquote:
                    self.current_para = self.doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in self.doc.styles] else self.doc.add_paragraph()
                else:
                    self.current_para = self.doc.add_paragraph()

            run = self.current_para.add_run(text)
            if self.bold:
                run.bold = True
            if self.italic:
                run.italic = True
            if self.underline:
                run.underline = True
            if self.strike:
                run.font.strike = True

    parser = HTMLToDocx(docx_doc)
    parser.feed(html_content)

    # Save to bytes
    buffer = BytesIO()
    docx_doc.save(buffer)
    buffer.seek(0)

    filename = re.sub(r'[^\w\s\-]', '', doc.get("title", "document")).strip() or "document"

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'}
    )


# ── Cross-Workspace Linking ──

@router.post("/{doc_id}/link-workspace")
async def link_document_to_workspace(
    doc_id: str,
    body: dict,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Link a document to an additional workspace."""
    user_id = _get_user_id(credentials)
    workspace_id = body.get("workspace_id")
    if not workspace_id:
        raise HTTPException(400, "workspace_id required")
    doc = await db.documents.find_one({"id": doc_id, "user_id": user_id}, {"_id": 0, "linked_workspaces": 1})
    if doc is None:
        raise HTTPException(404, "Document not found")
    linked = doc.get("linked_workspaces", [])
    if workspace_id in linked:
        return {"status": "already_linked"}
    await db.documents.update_one(
        {"id": doc_id},
        {"$addToSet": {"linked_workspaces": workspace_id}}
    )
    return {"status": "linked", "workspace_id": workspace_id}


@router.delete("/{doc_id}/unlink-workspace/{workspace_id}")
async def unlink_document_from_workspace(
    doc_id: str,
    workspace_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Remove a document's link to a workspace."""
    user_id = _get_user_id(credentials)
    result = await db.documents.update_one(
        {"id": doc_id, "user_id": user_id},
        {"$pull": {"linked_workspaces": workspace_id}}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Document not found")
    return {"status": "unlinked"}
