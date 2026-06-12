"""
AI Chat - File extraction and generation utilities.
Handles reading uploaded files (PDF, Excel, images, DOCX) and generating files (PDF, DOCX, XLSX).
"""
import io
import re
import logging
import base64 as b64

logger = logging.getLogger(__name__)


# ============== File Extraction ==============

def extract_pdf_text(file_bytes: bytes, max_chars: int = 8000) -> str:
    """Extract text from a PDF file."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages[:20]):
                page_text = page.extract_text() or ""
                text_parts.append(f"[Page {i+1}]\n{page_text}")
        full_text = "\n\n".join(text_parts)
        if len(full_text) > max_chars:
            return full_text[:max_chars] + "\n[TRUNCATED — document exceeds extraction limit]"
        return full_text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"[Error reading PDF: {str(e)}]"


def extract_excel_data(file_bytes: bytes, max_chars: int = 6000) -> str:
    """Extract data from an Excel/CSV file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        result = []
        for sheet_name in wb.sheetnames[:5]:
            ws = wb[sheet_name]
            result.append(f"[Sheet: {sheet_name}]")
            rows = []
            for row in ws.iter_rows(max_row=50, values_only=True):
                rows.append(" | ".join([str(c) if c is not None else "" for c in row]))
            result.append("\n".join(rows))
        text = "\n\n".join(result)
        if len(text) > max_chars:
            return text[:max_chars] + "\n[TRUNCATED]"
        return text
    except Exception as e:
        logger.error(f"Excel extraction error: {e}")
        return f"[Error reading spreadsheet: {str(e)}]"


def encode_image_base64(file_bytes: bytes) -> str:
    """Encode image bytes to base64 string."""
    return b64.b64encode(file_bytes).decode("utf-8")


def extract_docx_text(file_bytes: bytes, max_chars: int = 8000) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs[:100]])
        if len(text) > max_chars:
            return text[:max_chars] + "\n[TRUNCATED]"
        return text
    except Exception as e:
        return f"[Error reading DOCX: {str(e)}]"


async def extract_file_content(attachment: dict, get_object_fn) -> tuple:
    """Extract content from an uploaded file.
    Returns (text_description, image_data_url_or_None).
    """
    try:
        file_id = attachment.get("file_id")
        filename = attachment.get("filename", "")
        content_type = attachment.get("content_type", "")

        if not file_id:
            return f"[File: {filename}]", None

        file_bytes, _ = get_object_fn(f"ai-chat-files/{file_id}")

        if content_type.startswith("image/"):
            img_b64 = encode_image_base64(file_bytes)
            return f"[Image: {filename}]", f"data:{content_type};base64,{img_b64}"

        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            text = extract_pdf_text(file_bytes)
            return f"[PDF: {filename}]\n{text}", None

        ext = filename.lower().split(".")[-1] if "." in filename else ""
        if content_type in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel") or ext in ("xlsx", "xls"):
            text = extract_excel_data(file_bytes)
            return f"[Spreadsheet: {filename}]\n{text}", None

        if content_type.startswith("text/") or ext in ("txt", "csv", "json", "md", "py", "js", "ts", "html", "css"):
            text = file_bytes.decode("utf-8", errors="replace")[:8000]
            return f"[File: {filename}]\n{text}", None

        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == "docx":
            text = extract_docx_text(file_bytes)
            return f"[Document: {filename}]\n{text}", None

        return f"[Attached file: {filename} ({content_type})]", None
    except Exception as e:
        logger.error(f"File extraction error: {e}")
        return f"[Error reading file: {str(e)}]", None


# ============== File Generation ==============

def generate_pdf_from_markdown(text: str) -> bytes:
    """Convert markdown-like text to a PDF."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CustomH1', parent=styles['Heading1'], fontSize=18, spaceAfter=12))
    styles.add(ParagraphStyle(name='CustomH2', parent=styles['Heading2'], fontSize=14, spaceAfter=8))
    styles.add(ParagraphStyle(name='CustomBody', parent=styles['Normal'], fontSize=11, leading=16))

    story = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
        line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', line)
        line = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', line)

        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles['CustomH1']))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles['CustomH2']))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles['Heading3']))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"• {line[2:]}", styles['CustomBody']))
        else:
            story.append(Paragraph(line, styles['CustomBody']))

    doc.build(story)
    return buffer.getvalue()


def generate_docx_from_markdown(text: str) -> bytes:
    """Convert markdown-like text to a DOCX."""
    from docx import Document

    doc = Document()
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'`(.+?)`', r'\1', clean)

        if line.startswith("# "):
            doc.add_heading(clean[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(clean[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(clean[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(clean[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', clean), style='List Number')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_xlsx_from_text(text: str) -> bytes:
    """Extract table-like data from text and create an XLSX."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Generated Data"

    lines = text.split("\n")
    row_num = 1
    for line in lines:
        line = line.strip()
        if not line or line.startswith("---") or re.match(r'^[\|\-\s:]+$', line):
            continue
        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            for col_num, cell in enumerate(cells, 1):
                ws.cell(row=row_num, column=col_num, value=cell)
            row_num += 1
        elif line.startswith("- ") or line.startswith("* "):
            ws.cell(row=row_num, column=1, value=line[2:])
            row_num += 1
        elif not line.startswith("#") and not line.startswith("*"):
            ws.cell(row=row_num, column=1, value=line)
            row_num += 1

    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ============== Chart Generation ==============

def generate_pie_chart(data: dict) -> bytes:
    """Generate a pie chart PNG from JSON data."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    title = data.get("title", "Pie Chart")
    labels = data.get("labels", [])
    values = data.get("values", [])
    colors = data.get("colors", None)

    fig, ax = plt.subplots(figsize=(8, 6), facecolor='#0f172a')
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, colors=colors,
        autopct='%1.1f%%', startangle=90, pctdistance=0.82,
        wedgeprops=dict(width=0.5, edgecolor='#1e293b', linewidth=2),
    )
    for t in texts:
        t.set_color('white')
        t.set_fontsize(11)
    for t in autotexts:
        t.set_color('white')
        t.set_fontsize(10)
        t.set_fontweight('bold')
    ax.set_title(title, color='white', fontsize=16, fontweight='bold', pad=20)
    ax.legend(labels, loc='lower center', bbox_to_anchor=(0.5, -0.08), ncol=min(len(labels), 4),
              facecolor='#1e293b', edgecolor='#334155', labelcolor='white', fontsize=9)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return buf.getvalue()


def generate_bar_chart(data: dict) -> bytes:
    """Generate a bar chart PNG from JSON data."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    title = data.get("title", "Bar Chart")
    labels = data.get("labels", [])
    values = data.get("values", [])
    colors = data.get("colors", None)

    if not colors:
        default_palette = ['#7c3aed', '#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#06b6d4', '#f97316']
        colors = [default_palette[i % len(default_palette)] for i in range(len(labels))]

    fig, ax = plt.subplots(figsize=(10, 6), facecolor='#0f172a')
    ax.set_facecolor('#0f172a')
    bars = ax.bar(labels, values, color=colors, edgecolor='#1e293b', linewidth=1, width=0.6, zorder=3)

    # Add value labels on bars
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(values) * 0.02,
                str(val), ha='center', va='bottom', color='white', fontsize=10, fontweight='bold')

    ax.set_title(title, color='white', fontsize=16, fontweight='bold', pad=15)
    ax.tick_params(axis='x', colors='#94a3b8', labelsize=10)
    ax.tick_params(axis='y', colors='#94a3b8', labelsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#334155')
    ax.spines['bottom'].set_color('#334155')
    ax.grid(axis='y', color='#1e293b', linewidth=0.5, zorder=0)
    ax.set_axisbelow(True)
    plt.xticks(rotation=30 if len(labels) > 5 else 0, ha='right' if len(labels) > 5 else 'center')
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return buf.getvalue()


def generate_line_chart(data: dict) -> bytes:
    """Generate a line chart PNG from JSON data."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    title = data.get("title", "Line Chart")
    labels = data.get("labels", [])
    datasets = data.get("datasets", [])
    default_palette = ['#7c3aed', '#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#06b6d4', '#f97316', '#8b5cf6']

    # Support single-series shortcut
    if not datasets and data.get("values"):
        datasets = [{"name": data.get("series_name", "Value"), "values": data["values"]}]

    fig, ax = plt.subplots(figsize=(10, 6), facecolor='#0f172a')
    ax.set_facecolor('#0f172a')

    for i, ds in enumerate(datasets):
        color = (ds.get("color") or default_palette[i % len(default_palette)])
        values = ds.get("values", [])
        ax.plot(labels[:len(values)], values, color=color, linewidth=2.5, marker='o',
                markersize=6, markerfacecolor='white', markeredgecolor=color, markeredgewidth=2,
                label=ds.get("name", f"Series {i+1}"), zorder=3)
        ax.fill_between(labels[:len(values)], values, alpha=0.08, color=color, zorder=2)

    ax.set_title(title, color='white', fontsize=16, fontweight='bold', pad=15)
    ax.tick_params(axis='x', colors='#94a3b8', labelsize=9)
    ax.tick_params(axis='y', colors='#94a3b8', labelsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#334155')
    ax.spines['bottom'].set_color('#334155')
    ax.grid(axis='y', color='#1e293b', linewidth=0.5, zorder=0)
    ax.set_axisbelow(True)
    if len(datasets) > 1:
        ax.legend(facecolor='#1e293b', edgecolor='#334155', labelcolor='white', fontsize=9)
    plt.xticks(rotation=30 if len(labels) > 6 else 0, ha='right' if len(labels) > 6 else 'center')
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return buf.getvalue()


def generate_stacked_bar_chart(data: dict) -> bytes:
    """Generate a stacked bar chart PNG from JSON data."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    title = data.get("title", "Stacked Bar Chart")
    labels = data.get("labels", [])
    datasets = data.get("datasets", [])
    default_palette = ['#7c3aed', '#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#06b6d4', '#f97316', '#8b5cf6']

    fig, ax = plt.subplots(figsize=(10, 6), facecolor='#0f172a')
    ax.set_facecolor('#0f172a')

    x = np.arange(len(labels))
    bottoms = np.zeros(len(labels))
    bar_width = 0.5

    for i, ds in enumerate(datasets):
        color = ds.get("color", default_palette[i % len(default_palette)])
        values = np.array(ds.get("values", [0] * len(labels))[:len(labels)], dtype=float)
        ax.bar(x, values, bar_width, bottom=bottoms, color=color, edgecolor='#1e293b',
               linewidth=0.5, label=ds.get("name", f"Series {i+1}"), zorder=3)
        bottoms += values

    ax.set_title(title, color='white', fontsize=16, fontweight='bold', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, color='#94a3b8', fontsize=10)
    ax.tick_params(axis='y', colors='#94a3b8', labelsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#334155')
    ax.spines['bottom'].set_color('#334155')
    ax.grid(axis='y', color='#1e293b', linewidth=0.5, zorder=0)
    ax.set_axisbelow(True)
    ax.legend(facecolor='#1e293b', edgecolor='#334155', labelcolor='white', fontsize=9,
              loc='upper left')
    plt.xticks(rotation=30 if len(labels) > 5 else 0, ha='right' if len(labels) > 5 else 'center')
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return buf.getvalue()


def generate_radar_chart(data: dict) -> bytes:
    """Generate a radar/spider chart PNG from JSON data."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    title = data.get("title", "Radar Chart")
    labels = data.get("labels", [])
    datasets = data.get("datasets", [])
    default_palette = ['#7c3aed', '#3b82f6', '#10b981', '#f59e0b', '#ef4444']

    if not datasets and data.get("values"):
        datasets = [{"name": data.get("series_name", "Value"), "values": data["values"]}]

    n = len(labels)
    if n < 3:
        raise ValueError("Radar chart needs at least 3 axes")

    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True), facecolor='#0f172a')
    ax.set_facecolor('#0f172a')

    for i, ds in enumerate(datasets):
        color = ds.get("color", default_palette[i % len(default_palette)])
        values = ds.get("values", [0] * n)[:n]
        values += values[:1]
        ax.plot(angles, values, color=color, linewidth=2.5, label=ds.get("name", f"Series {i+1}"))
        ax.fill(angles, values, color=color, alpha=0.15)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, color='#e2e8f0', fontsize=10)
    ax.tick_params(axis='y', colors='#475569', labelsize=8)
    ax.set_rlabel_position(30)
    ax.spines['polar'].set_color('#334155')
    ax.grid(color='#1e293b', linewidth=0.5)
    ax.set_title(title, color='white', fontsize=16, fontweight='bold', pad=25, y=1.08)
    if len(datasets) > 1:
        ax.legend(facecolor='#1e293b', edgecolor='#334155', labelcolor='white', fontsize=9,
                  loc='lower right', bbox_to_anchor=(1.15, -0.05))
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return buf.getvalue()
