"""
AI routes - TTS, transcription, chat, video generation.
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Depends, Query
from fastapi.responses import StreamingResponse, Response
from typing import List, Optional
from pydantic import BaseModel
from datetime import datetime, timezone
import os
import io
import base64

from config import logger, db

router = APIRouter(tags=["AI"])

# OpenAI configuration
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')


# ============== Models ==============

class TranscriptAnalyzeRequest(BaseModel):
    transcript: str
    analysis_type: str = "summary"

class TTSRequest(BaseModel):
    text: str
    voice: str = "alloy"
    speed: float = 1.0

class AIChatMessage(BaseModel):
    role: str
    content: str

class AIChatRequest(BaseModel):
    message: str
    conversation_history: List[dict] = []
    system_prompt: Optional[str] = None
    user_id: Optional[str] = None  # For entitlement checking


# ============== Routes ==============

@router.post("/transcripts/analyze")
async def analyze_transcript(request: TranscriptAnalyzeRequest):
    """Analyze a transcript using AI"""
    try:
        api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        from llm_client import chat_completion
        
        prompts = {
            "summary": "Provide a concise summary of this transcript:",
            "action_items": "Extract all action items from this transcript:",
            "key_points": "List the key points discussed in this transcript:",
            "sentiment": "Analyze the overall sentiment of this transcript:",
            "questions": "List any questions raised in this transcript:"
        }
        
        prompt = prompts.get(request.analysis_type, prompts["summary"])
        
        response = chat_completion(
            api_key=api_key,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that analyzes meeting transcripts."},
                {"role": "user", "content": f"{prompt}\n\n{request.transcript}"}
            ],
            model="gpt-4o"
        )
        
        return {
            "success": True,
            "analysis_type": request.analysis_type,
            "result": response
        }
    except Exception as e:
        logger.error(f"Transcript analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/tts/voices")
async def get_tts_voices():
    """Get available TTS voices"""
    voices = [
        {"id": "alloy", "name": "Alloy", "description": "Neutral and balanced"},
        {"id": "echo", "name": "Echo", "description": "Male, warm"},
        {"id": "fable", "name": "Fable", "description": "British accent"},
        {"id": "onyx", "name": "Onyx", "description": "Male, deep"},
        {"id": "nova", "name": "Nova", "description": "Female, friendly"},
        {"id": "shimmer", "name": "Shimmer", "description": "Female, soft"}
    ]
    return {"voices": voices}


@router.post("/tts/generate")
async def generate_tts(request: TTSRequest):
    """Generate text-to-speech audio"""
    try:
        api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="TTS service not configured")
        
        from llm_client import text_to_speech
        
        response = text_to_speech(
            text=request.text,
            voice=request.voice,
            model="tts-1",
            api_key=api_key,
        )
        audio_bytes = response.content
        
        return StreamingResponse(
            iter([audio_bytes]),
            media_type="audio/mpeg",
            headers={"Content-Disposition": "attachment; filename=speech.mp3"}
        )
    except Exception as e:
        logger.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/tts/generate-base64")
async def generate_tts_base64(request: TTSRequest):
    """Generate TTS and return as base64"""
    try:
        api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="TTS service not configured")
        
        from llm_client import text_to_speech
        
        response = text_to_speech(
            text=request.text,
            voice=request.voice,
            model="tts-1",
            api_key=api_key,
        )
        audio_bytes = response.content
        
        audio_base64 = base64.b64encode(audio_bytes).decode('utf-8')
        
        return {
            "success": True,
            "audio": audio_base64,
            "mime_type": "audio/mpeg"
        }
    except Exception as e:
        logger.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/ai/chat")
async def ai_chat(request: AIChatRequest):
    """Chat with AI assistant"""
    try:
        # Check entitlements if user_id provided
        if request.user_id:
            from routes.entitlements import check_entitlement, record_usage
            check = await check_entitlement(request.user_id, "ai_chat", 1)
            if not check.allowed:
                raise HTTPException(
                    status_code=403,
                    detail={
                        "code": "AI_CHAT_LIMIT_REACHED",
                        "message": check.message,
                        "upgrade_url": "/pricing"
                    }
                )
        
        api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        from llm_client import chat_completion
        import uuid
        
        # Build system message
        system_message = request.system_prompt or "You are Munal AI, a helpful assistant for a meeting and collaboration platform. Be concise and helpful."
        
        # Build messages for API
        messages = [{"role": "system", "content": system_message}]
        for msg in request.conversation_history:
            messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
        messages.append({"role": "user", "content": request.message})
        
        # Send to API
        result = chat_completion(
            messages=messages,
            model="gpt-4o",
            api_key=api_key,
        )
        response = result.choices[0].message.content
        
        # Record usage after successful response
        if request.user_id:
            await record_usage(request.user_id, "ai_chat", 1, {"message_length": len(request.message)})
        
        return {
            "success": True,
            "response": response
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/ai/chat/stream")
async def ai_chat_stream(request: AIChatRequest):
    """Stream chat response from AI - falls back to non-streaming"""
    # For now, use non-streaming endpoint since the library doesn't expose stream directly
    return await ai_chat(request)


@router.post("/ai/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    language: str = Form(default="en")
):
    """Transcribe audio file using OpenAI Whisper - uses platform API key"""
    try:
        api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="Transcription service not configured")
        
        # Validate file size (25MB limit)
        contents = await file.read()
        if len(contents) > 25 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File size exceeds 25MB limit")
        
        # Validate file type
        valid_extensions = ['.mp3', '.mp4', '.mpeg', '.mpga', '.m4a', '.wav', '.webm', '.ogg']
        file_ext = '.' + file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        if file_ext not in valid_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Invalid file format. Supported: {', '.join(valid_extensions)}"
            )
        
        from llm_client import speech_to_text
        import io
        
        # Create file-like object from contents
        audio_file = io.BytesIO(contents)
        audio_file.name = file.filename
        
        # Transcribe
        response = speech_to_text(audio_file=audio_file, api_key=api_key)
        
        # Build response
        result = {
            "success": True,
            "text": response.text,
            "language": getattr(response, 'language', language),
            "duration": getattr(response, 'duration', None),
            "segments": []
        }
        
        # Add segments if available
        if hasattr(response, 'segments') and response.segments:
            result["segments"] = [
                {
                    "start": seg.start if hasattr(seg, 'start') else seg.get('start'),
                    "end": seg.end if hasattr(seg, 'end') else seg.get('end'),
                    "text": seg.text if hasattr(seg, 'text') else seg.get('text')
                }
                for seg in response.segments
            ]
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Transcription error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ai/transcribe/status")
async def get_transcription_status():
    """Check if transcription service is available (platform has API key configured)"""
    api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
    return {
        "available": bool(api_key),
        "provider": "openai_whisper"
    }


# ============== Recording Transcription ==============

import uuid
from motor.motor_asyncio import AsyncIOMotorGridFSBucket

class RecordingTranscriptionRequest(BaseModel):
    file_id: str
    user_id: str
    file_name: str

@router.post("/ai/transcribe/recording")
async def transcribe_recording(request: RecordingTranscriptionRequest):
    """Transcribe a stored recording from GridFS"""
    try:
        api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="Transcription service not configured")
        
        # Get the file from GridFS
        fs = AsyncIOMotorGridFSBucket(db, bucket_name="chat_files")
        
        # Find the file metadata
        file_doc = await db.chat_files.find_one({"id": request.file_id})
        if not file_doc:
            raise HTTPException(status_code=404, detail="Recording not found")
        
        grid_id = file_doc.get("grid_id")
        if not grid_id:
            raise HTTPException(status_code=404, detail="Recording file not found in storage")
        
        from bson import ObjectId
        import io
        import tempfile
        import subprocess
        import os
        
        # Download the file from GridFS
        file_data = io.BytesIO()
        await fs.download_to_stream(ObjectId(grid_id), file_data)
        file_data.seek(0)
        
        # Check file size (25MB limit for Whisper)
        file_size = file_data.getbuffer().nbytes
        audio_file = file_data
        temp_files = []
        
        if file_size > 25 * 1024 * 1024:
            # Large file - extract audio using ffmpeg
            logger.info(f"Large file detected ({file_size / (1024*1024):.1f}MB), extracting audio with ffmpeg")
            
            try:
                # Save video to temp file
                with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as video_temp:
                    video_temp.write(file_data.read())
                    video_temp_path = video_temp.name
                    temp_files.append(video_temp_path)
                
                # Create temp file for audio output
                audio_temp_fd, audio_temp_path = tempfile.mkstemp(suffix='.mp3')
                os.close(audio_temp_fd)
                temp_files.append(audio_temp_path)
                
                # Extract audio using ffmpeg (compress to mp3 at lower bitrate)
                ffmpeg_cmd = [
                    'ffmpeg', '-y',
                    '-i', video_temp_path,
                    '-vn',  # No video
                    '-acodec', 'libmp3lame',
                    '-ab', '64k',  # Lower bitrate for smaller file
                    '-ar', '16000',  # 16kHz sample rate (good for speech)
                    '-ac', '1',  # Mono
                    audio_temp_path
                ]
                
                result = subprocess.run(
                    ffmpeg_cmd, 
                    capture_output=True, 
                    text=True,
                    timeout=300  # 5 minute timeout
                )
                
                if result.returncode != 0:
                    logger.error(f"FFmpeg error: {result.stderr}")
                    raise Exception(f"Audio extraction failed: {result.stderr[:200]}")
                
                # Check extracted audio size
                extracted_size = os.path.getsize(audio_temp_path)
                logger.info(f"Extracted audio size: {extracted_size / (1024*1024):.1f}MB")
                
                if extracted_size > 25 * 1024 * 1024:
                    raise HTTPException(
                        status_code=400,
                        detail="Recording audio too long for transcription. Please split into shorter segments."
                    )
                
                # Read the extracted audio
                audio_file = io.BytesIO()
                with open(audio_temp_path, 'rb') as f:
                    audio_file.write(f.read())
                audio_file.seek(0)
                audio_file.name = 'audio.mp3'
                
            except subprocess.TimeoutExpired:
                raise HTTPException(status_code=500, detail="Audio extraction timed out")
            except Exception as e:
                logger.error(f"FFmpeg extraction error: {e}")
                raise HTTPException(status_code=500, detail=f"Audio extraction failed: {str(e)}")
            finally:
                # Cleanup temp files
                for temp_path in temp_files:
                    try:
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)
                    except Exception:
                        pass
        else:
            # Prepare file for transcription
            audio_file.name = request.file_name or "recording.webm"
        
        from llm_client import speech_to_text
        
        # Initialize STT
        response = speech_to_text(audio_file=audio_file, api_key=api_key)
        
        # Build transcript data
        transcript_id = str(uuid.uuid4())
        segments = []
        
        if hasattr(response, 'segments') and response.segments:
            segments = [
                {
                    "start": seg.start if hasattr(seg, 'start') else seg.get('start'),
                    "end": seg.end if hasattr(seg, 'end') else seg.get('end'),
                    "text": seg.text if hasattr(seg, 'text') else seg.get('text')
                }
                for seg in response.segments
            ]
        
        # Store transcript in database
        transcript_doc = {
            "id": transcript_id,
            "file_id": request.file_id,
            "user_id": request.user_id,
            "file_name": request.file_name,
            "text": response.text,
            "language": getattr(response, 'language', 'en'),
            "duration": getattr(response, 'duration', None),
            "segments": segments,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "status": "completed"
        }
        
        await db.recording_transcripts.insert_one(transcript_doc)
        
        # Update the file record with transcript reference
        await db.chat_files.update_one(
            {"id": request.file_id},
            {"$set": {"transcript_id": transcript_id, "has_transcript": True}}
        )
        
        return {
            "success": True,
            "transcript_id": transcript_id,
            "text": response.text,
            "duration": getattr(response, 'duration', None),
            "segments": segments
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Recording transcription error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ai/transcribe/recording/{file_id}")
async def get_recording_transcript(file_id: str):
    """Get transcript for a recording"""
    try:
        transcript = await db.recording_transcripts.find_one(
            {"file_id": file_id},
            {"_id": 0}
        )
        
        if not transcript:
            return {"success": False, "message": "No transcript found for this recording"}
        
        return {"success": True, "transcript": transcript}
        
    except Exception as e:
        logger.error(f"Get transcript error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ai/transcripts/user/{user_id}")
async def get_user_transcripts(user_id: str, limit: int = 50):
    """Get all transcripts for a user"""
    try:
        transcripts = await db.recording_transcripts.find(
            {"user_id": user_id},
            {"_id": 0}
        ).sort("created_at", -1).limit(limit).to_list(limit)
        
        return {"success": True, "transcripts": transcripts, "count": len(transcripts)}
        
    except Exception as e:
        logger.error(f"Get user transcripts error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/ai/transcripts/{transcript_id}")
async def delete_transcript(transcript_id: str):
    """Delete a transcript"""
    try:
        result = await db.recording_transcripts.delete_one({"id": transcript_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Transcript not found")
        
        return {"success": True, "message": "Transcript deleted"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete transcript error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============== Video Generation (Sora 2) ==============

import asyncio
from concurrent.futures import ThreadPoolExecutor

# Store for video generation jobs
video_jobs = {}
video_executor = ThreadPoolExecutor(max_workers=2)

class VideoGenerationRequest(BaseModel):
    prompt: str
    size: str = "1280x720"  # 1280x720, 1792x1024, 1024x1792, 1024x1024
    duration: int = 12  # Base: 4, 8, 12. Extended: 24, 36, 48, 60 (multi-clip)
    model: str = "sora-2"  # sora-2 or sora-2-pro
    voice: str = "nova"  # alloy, echo, fable, onyx, nova, shimmer


async def get_video_api_key():
    """Get the Emergent LLM key for video generation"""
    return EMERGENT_LLM_KEY


def generate_video_sync(job_id: str, prompt: str, model: str, size: str, duration: int, api_key: str, voice: str = "nova"):
    """Synchronous video generation using Emergent's OpenAIVideoGeneration (runs in thread pool)"""
    import os
    from emergentintegrations.llm.openai.video_generation import OpenAIVideoGeneration

    try:
        video_jobs[job_id] = {"status": "generating", "progress": 10, "message": "Starting video generation..."}

        base_durations = [4, 8, 12]
        extended_durations = [24, 36, 48, 60]

        def generate_single_clip(clip_prompt, dur, clip_idx=0, total_clips=1):
            """Generate a single video clip using Emergent"""
            video_gen = OpenAIVideoGeneration(api_key=api_key)
            video_bytes = video_gen.text_to_video(
                prompt=clip_prompt, model=model, size=size,
                duration=dur, max_wait_time=600
            )
            if not video_bytes:
                raise Exception(f"Clip {clip_idx+1} returned no data")
            clip_path = f"/tmp/clip_{job_id}_{clip_idx}.mp4"
            video_gen.save_video(video_bytes, clip_path)
            return clip_path, video_bytes

        if duration in extended_durations:
            num_clips = duration // 12
            clip_duration = 12
            clip_paths = []

            for i in range(num_clips):
                video_jobs[job_id] = {
                    "status": "generating",
                    "progress": 10 + (i * 70 // num_clips),
                    "message": f"Generating clip {i+1}/{num_clips}..."
                }
                clip_prompt = f"Continuation of scene: {prompt}" if i > 0 else prompt
                clip_path, _ = generate_single_clip(clip_prompt, clip_duration, i, num_clips)
                clip_paths.append(clip_path)

            video_jobs[job_id] = {"status": "generating", "progress": 90, "message": "Stitching clips..."}
            output_path = f"/tmp/video_{job_id}.mp4"

            if not stitch_videos_with_ffmpeg(clip_paths, output_path):
                video_jobs[job_id] = {"status": "failed", "error": "Failed to stitch clips"}
                return

            with open(output_path, 'rb') as f:
                video_bytes = f.read()
            os.remove(output_path)
            for cp in clip_paths:
                try: os.remove(cp)
                except: pass
        else:
            # Standard duration (4, 8, or 12)
            # Snap to valid
            valid = [4, 8, 12]
            duration = min(valid, key=lambda x: abs(x - duration))

            video_jobs[job_id] = {"status": "generating", "progress": 20, "message": "Creating video..."}

            video_jobs[job_id] = {"status": "generating", "progress": 30, "message": "Generating video (this may take a few minutes)..."}

            max_retries = 3
            video_bytes = None
            for attempt in range(max_retries):
                try:
                    # IMPORTANT: Always create a NEW instance per attempt
                    video_gen = OpenAIVideoGeneration(api_key=api_key)
                    wait_time = 900 if duration >= 12 else 600
                    video_bytes = video_gen.text_to_video(
                        prompt=prompt, model=model, size=size,
                        duration=duration, max_wait_time=wait_time
                    )
                    if video_bytes and len(video_bytes) > 100:
                        break
                    logger.warning(f"Video job {job_id} attempt {attempt+1} returned empty data")
                    video_bytes = None
                except Exception as e:
                    logger.warning(f"Video job {job_id} attempt {attempt+1} failed: {e}")
                    video_bytes = None
                if attempt < max_retries - 1:
                    import time
                    time.sleep(10 * (attempt + 1))

        if not video_bytes:
            video_jobs[job_id] = {"status": "failed", "error": "No video returned"}
            return

        video_base64 = base64.b64encode(video_bytes).decode('utf-8')
        video_jobs[job_id] = {
            "status": "completed",
            "progress": 100,
            "video_base64": video_base64,
            "size": size,
            "duration": duration,
            "voice": voice,
            "file_size": len(video_bytes)
        }
        logger.info(f"Video job {job_id} completed: {len(video_bytes)} bytes")
        
    except Exception as e:
        logger.error(f"Video job {job_id} failed: {e}")
        video_jobs[job_id] = {"status": "failed", "error": str(e)}


def stitch_videos_with_ffmpeg(video_paths: list, output_path: str) -> bool:
    """Stitch multiple videos together using ffmpeg"""
    import subprocess
    import os
    
    # Create a file list for ffmpeg concat
    list_file = f"/tmp/video_list_{os.path.basename(output_path)}.txt"
    with open(list_file, 'w') as f:
        for path in video_paths:
            f.write(f"file '{path}'\n")
    
    try:
        # Use ffmpeg concat demuxer for seamless stitching
        cmd = [
            'ffmpeg', '-y', '-f', 'concat', '-safe', '0',
            '-i', list_file,
            '-c', 'copy',  # Copy without re-encoding for speed
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, timeout=120)
        
        # Cleanup temp files
        os.remove(list_file)
        for path in video_paths:
            if os.path.exists(path):
                os.remove(path)
        
        return result.returncode == 0
    except Exception as e:
        logger.error(f"FFmpeg stitching error: {e}")
        return False


@router.post("/ai/video/generate")
async def generate_video(request: VideoGenerationRequest):
    """Start async video generation job. Returns job_id immediately for polling."""
    try:
        # Get API key from admin settings or fallback to env
        api_key = await get_video_api_key()
        
        if not api_key:
            raise HTTPException(
                status_code=500, 
                detail="Video generation not configured. Please contact administrator."
            )
        
        # Validate parameters
        valid_sizes = ["1280x720", "1792x1024", "1024x1792", "1024x1024"]
        base_durations = [4, 8, 12]
        extended_durations = [24, 36, 48, 60]
        all_valid_durations = base_durations + extended_durations
        valid_models = ["sora-2", "sora-2-pro"]
        
        if request.size not in valid_sizes:
            raise HTTPException(status_code=400, detail=f"Invalid size. Must be one of: {valid_sizes}")
        if request.duration not in all_valid_durations:
            raise HTTPException(status_code=400, detail=f"Invalid duration. Must be one of: {all_valid_durations}")
        if request.model not in valid_models:
            raise HTTPException(status_code=400, detail=f"Invalid model. Must be one of: {valid_models}")
        
        import uuid
        job_id = str(uuid.uuid4())
        
        # Initialize job status
        video_jobs[job_id] = {"status": "queued", "progress": 0}
        
        # Start generation in background thread
        video_executor.submit(
            generate_video_sync,
            job_id,
            request.prompt,
            request.model,
            request.size,
            request.duration,
            api_key,
            request.voice
        )
        
        logger.info(f"Video job {job_id} started: prompt='{request.prompt[:50]}...', duration={request.duration}s, voice={request.voice}")
        
        return {
            "success": True,
            "job_id": job_id,
            "status": "queued",
            "message": "Video generation started. Poll /api/ai/video/job/{job_id} for status."
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Video generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ai/video/job/{job_id}")
async def get_video_job_status(job_id: str):
    """Poll video generation job status"""
    if job_id not in video_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = video_jobs[job_id]
    
    # If completed, include the video data
    if job.get("status") == "completed":
        return {
            "success": True,
            "status": "completed",
            "progress": 100,
            "video_base64": job.get("video_base64"),
            "size": job.get("size"),
            "duration": job.get("duration"),
            "file_size": job.get("file_size"),
            "mime_type": "video/mp4"
        }
    elif job.get("status") == "failed":
        return {
            "success": False,
            "status": "failed",
            "error": job.get("error", "Unknown error")
        }
    else:
        return {
            "success": True,
            "status": job.get("status", "processing"),
            "progress": job.get("progress", 0),
            "message": job.get("message", "Generating video...")
        }


@router.get("/ai/video/status")
async def get_video_generation_status():
    """Check if video generation service is available"""
    api_key = await get_video_api_key()
    return {
        "available": bool(api_key),
        "provider": "sora-2",
        "supported_sizes": ["1280x720", "1792x1024", "1024x1792", "1024x1024"],
        "supported_durations": {
            "base": [4, 8, 12],
            "extended": [24, 36, 48, 60],
            "scene_based": [60, 120, 180, 240, 300]
        },
        "supported_models": ["sora-2", "sora-2-pro"],
        "supported_voices": ["alloy", "echo", "fable", "onyx", "nova", "shimmer"],
        "max_scene_duration": 300,
    }


# ============== Scene-Based Video Pipeline ==============

class SceneSplitRequest(BaseModel):
    prompt: str
    target_duration: int = 180  # seconds (60-300)
    scene_length: int = 30  # seconds per scene (10-60)


class SceneGenerateRequest(BaseModel):
    scenes: list  # [{prompt, duration}]
    model: str = "sora-2"
    size: str = "1280x720"
    voice: str = "nova"


@router.post("/ai/video/split-scenes")
async def split_prompt_into_scenes(req: SceneSplitRequest):
    """Use AI to split a prompt into multiple scenes for a long video"""
    api_key = await get_video_api_key()
    if not api_key:
        raise HTTPException(status_code=500, detail="Video service not configured")

    target = max(60, min(300, req.target_duration))
    scene_len = max(4, min(12, req.scene_length))
    # Snap scene_len to valid Sora 2 values
    valid_lens = [4, 8, 12]
    scene_len = min(valid_lens, key=lambda x: abs(x - scene_len))
    num_scenes = max(2, min(10, target // scene_len))

    from emergentintegrations.llm.chat import LlmChat, UserMessage
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY or api_key,
        session_id=f"scene-split-{uuid.uuid4()}",
        system_message="You are a professional video director. Split prompts into distinct visual scenes for AI video generation. Return ONLY valid JSON."
    ).with_model("openai", "gpt-5.5")

    split_prompt = f"""Split this video concept into exactly {num_scenes} scenes, each {scene_len} seconds long.

Video concept: "{req.prompt}"
Total duration: {target} seconds

Return a JSON array of scenes. Each scene should have:
- "scene_number": integer
- "prompt": detailed visual description for AI video generation (camera angles, lighting, movement, style)
- "duration": {scene_len}
- "transition": how this scene connects to the next ("cut", "fade", "dissolve")

Rules:
- Each prompt should be self-contained but flow naturally from the previous scene
- Include specific visual details: camera movement, lighting, colors, mood
- Make each scene visually distinct but part of a coherent narrative
- Start the first scene with an establishing shot
- End the last scene with a closing shot

Return ONLY the JSON array, no markdown or explanation."""

    response = await chat.send_message(UserMessage(text=split_prompt))
    response_text = response if isinstance(response, str) else getattr(response, "text", str(response))

    # Parse JSON from response
    import json
    try:
        # Strip markdown if present
        text = response_text.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        scenes = json.loads(text)
    except json.JSONDecodeError:
        # Try to extract JSON array
        start = text.find("[")
        end = text.rfind("]") + 1
        if start >= 0 and end > start:
            scenes = json.loads(text[start:end])
        else:
            raise HTTPException(status_code=500, detail="Failed to parse scenes from AI response")

    return {
        "scenes": scenes,
        "total_duration": sum(s.get("duration", scene_len) for s in scenes),
        "scene_count": len(scenes),
    }


def generate_scenes_parallel(job_id: str, scenes: list, model: str, size: str, api_key: str, voice: str = "nova"):
    """Generate multiple scenes in parallel using Emergent and stitch them together"""
    import os
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from emergentintegrations.llm.openai.video_generation import OpenAIVideoGeneration

    try:
        video_jobs[job_id] = {"status": "generating", "progress": 5, "message": "Starting scene generation...", "scenes_total": len(scenes), "scenes_done": 0}

        def generate_single_scene(scene_idx, scene):
            """Generate a single scene clip using Emergent with retries"""
            prompt_text = scene.get("prompt", "")
            dur = min(12, max(4, scene.get("duration", 8)))

            # Snap to valid duration
            valid = [4, 8, 12]
            dur = min(valid, key=lambda x: abs(x - dur))

            max_retries = 3
            for attempt in range(max_retries):
                try:
                    # IMPORTANT: Always create a NEW instance per attempt
                    video_gen = OpenAIVideoGeneration(api_key=api_key)
                    wait_time = 900 if dur >= 12 else 600
                    video_bytes = video_gen.text_to_video(
                        prompt=prompt_text, model=model, size=size,
                        duration=dur, max_wait_time=wait_time
                    )
                    if video_bytes and len(video_bytes) > 100:
                        clip_path = f"/tmp/scene_{job_id}_{scene_idx}.mp4"
                        video_gen.save_video(video_bytes, clip_path)
                        logger.info(f"Scene {scene_idx+1} generated: {len(video_bytes)} bytes")
                        return scene_idx, clip_path
                    else:
                        logger.warning(f"Scene {scene_idx+1} attempt {attempt+1} returned empty data")
                except Exception as e:
                    logger.warning(f"Scene {scene_idx+1} attempt {attempt+1} failed: {e}")

                if attempt < max_retries - 1:
                    import time
                    time.sleep(10 * (attempt + 1))  # Backoff: 10s, 20s

            raise Exception(f"Scene {scene_idx+1} failed after {max_retries} attempts")

        # Generate scenes in parallel (max 2 concurrent to avoid rate limits)
        clip_paths = [None] * len(scenes)
        scenes_done = 0

        with ThreadPoolExecutor(max_workers=2) as executor:
            futures = {executor.submit(generate_single_scene, i, s): i for i, s in enumerate(scenes)}
            for future in as_completed(futures):
                idx = futures[future]
                try:
                    scene_idx, path = future.result()
                    clip_paths[scene_idx] = path
                    scenes_done += 1
                    pct = 10 + int((scenes_done / len(scenes)) * 75)
                    video_jobs[job_id] = {
                        "status": "generating", "progress": pct,
                        "message": f"Scene {scenes_done}/{len(scenes)} complete",
                        "scenes_total": len(scenes), "scenes_done": scenes_done
                    }
                except Exception as e:
                    logger.error(f"Scene {idx} failed: {e}")
                    video_jobs[job_id] = {"status": "failed", "error": str(e)}
                    return

        # Verify all clips
        valid_clips = [p for p in clip_paths if p]
        if len(valid_clips) < len(scenes):
            video_jobs[job_id] = {"status": "failed", "error": f"Only {len(valid_clips)}/{len(scenes)} scenes generated"}
            return

        # Stitch
        video_jobs[job_id] = {"status": "generating", "progress": 90, "message": "Stitching scenes together...", "scenes_total": len(scenes), "scenes_done": len(scenes)}
        output_path = f"/tmp/final_{job_id}.mp4"

        if not stitch_videos_with_ffmpeg(valid_clips, output_path):
            video_jobs[job_id] = {"status": "failed", "error": "FFmpeg stitching failed"}
            return

        import os
        with open(output_path, 'rb') as f:
            video_bytes = f.read()
        os.remove(output_path)

        video_base64 = base64.b64encode(video_bytes).decode('utf-8')
        total_dur = sum(s.get("duration", 30) for s in scenes)
        video_jobs[job_id] = {
            "status": "completed", "progress": 100,
            "video_base64": video_base64, "size": size,
            "duration": total_dur, "voice": voice,
            "file_size": len(video_bytes),
            "scenes_total": len(scenes), "scenes_done": len(scenes),
        }
        logger.info(f"Scene-based video {job_id} completed: {len(scenes)} scenes, {len(video_bytes)} bytes")

    except Exception as e:
        logger.error(f"Scene-based video {job_id} failed: {e}")
        video_jobs[job_id] = {"status": "failed", "error": str(e)}


@router.post("/ai/video/generate-scenes")
async def generate_video_from_scenes(request: SceneGenerateRequest):
    """Generate a long video from multiple scenes (parallel generation + FFmpeg stitch)"""
    api_key = await get_video_api_key()
    if not api_key:
        raise HTTPException(status_code=500, detail="Video service not configured")

    if not request.scenes or len(request.scenes) < 2:
        raise HTTPException(status_code=400, detail="At least 2 scenes required")
    if len(request.scenes) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 scenes per video")

    job_id = str(uuid.uuid4())
    video_jobs[job_id] = {"status": "queued", "progress": 0, "scenes_total": len(request.scenes), "scenes_done": 0}

    video_executor.submit(
        generate_scenes_parallel,
        job_id, request.scenes, request.model, request.size, api_key, request.voice
    )

    total_dur = sum(s.get("duration", 30) for s in request.scenes)
    logger.info(f"Scene-based job {job_id} started: {len(request.scenes)} scenes, ~{total_dur}s total")

    return {
        "success": True,
        "job_id": job_id,
        "status": "queued",
        "scene_count": len(request.scenes),
        "estimated_duration": total_dur,
        "message": f"Generating {len(request.scenes)} scenes in parallel. Poll /api/ai/video/job/{job_id} for status."
    }


# ============== Video History ==============

class SaveVideoRequest(BaseModel):
    video_base64: str
    prompt: str
    duration: int
    size: str
    title: Optional[str] = None


@router.post("/ai/video/history")
async def save_video_to_history(request: SaveVideoRequest):
    """Save generated video to user's history"""
    try:
        # Calculate file size from base64
        video_bytes = base64.b64decode(request.video_base64)
        file_size = len(video_bytes)
        
        video_doc = {
            "title": request.title or f"Video - {request.duration}s",
            "prompt": request.prompt,
            "duration": request.duration,
            "size": request.size,
            "file_size": file_size,
            "video_base64": request.video_base64,
            "created_at": datetime.now(timezone.utc),
            "mime_type": "video/mp4"
        }
        
        result = await db.video_history.insert_one(video_doc)
        
        return {
            "success": True,
            "video_id": str(result.inserted_id),
            "message": "Video saved to history"
        }
    except Exception as e:
        logger.error(f"Error saving video to history: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ai/video/history")
async def get_video_history(limit: int = 20, skip: int = 0):
    """Get user's video history (without video data for listing)"""
    try:
        # Get videos without the large base64 data
        cursor = db.video_history.find(
            {},
            {"video_base64": 0}  # Exclude large video data
        ).sort("created_at", -1).skip(skip).limit(limit)
        
        videos = []
        async for doc in cursor:
            videos.append({
                "id": str(doc["_id"]),
                "title": doc.get("title", "Untitled"),
                "prompt": doc.get("prompt", ""),
                "duration": doc.get("duration"),
                "size": doc.get("size"),
                "file_size": doc.get("file_size"),
                "created_at": doc.get("created_at").isoformat() if doc.get("created_at") else None
            })
        
        total = await db.video_history.count_documents({})
        
        return {
            "videos": videos,
            "total": total,
            "limit": limit,
            "skip": skip
        }
    except Exception as e:
        logger.error(f"Error fetching video history: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/ai/video/history/{video_id}")
async def get_video_from_history(video_id: str):
    """Get a specific video from history (includes video data)"""
    try:
        from bson import ObjectId
        
        doc = await db.video_history.find_one({"_id": ObjectId(video_id)})
        
        if not doc:
            raise HTTPException(status_code=404, detail="Video not found")
        
        return {
            "id": str(doc["_id"]),
            "title": doc.get("title", "Untitled"),
            "prompt": doc.get("prompt", ""),
            "duration": doc.get("duration"),
            "size": doc.get("size"),
            "file_size": doc.get("file_size"),
            "video_base64": doc.get("video_base64"),
            "mime_type": doc.get("mime_type", "video/mp4"),
            "created_at": doc.get("created_at").isoformat() if doc.get("created_at") else None
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching video: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/ai/video/history/{video_id}")
async def delete_video_from_history(video_id: str):
    """Delete a video from history"""
    try:
        from bson import ObjectId
        
        result = await db.video_history.delete_one({"_id": ObjectId(video_id)})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Video not found")
        
        return {"success": True, "message": "Video deleted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting video: {e}")
        raise HTTPException(status_code=500, detail=str(e))



# ============== Meeting Auto-Transcribe + Insights ==============

class MeetingProcessRequest(BaseModel):
    meeting_id: str
    user_id: str
    meeting_title: Optional[str] = ""
    participants: Optional[List[str]] = []
    duration_seconds: Optional[int] = 0


@router.post("/ai/meeting/process")
async def process_meeting_audio(
    file: UploadFile = File(...),
    meeting_id: str = Form(...),
    user_id: str = Form(...),
    meeting_title: str = Form(default=""),
    participants: str = Form(default=""),
    duration_seconds: int = Form(default=0),
):
    """Auto-transcribe meeting audio and generate AI insights."""
    api_key = EMERGENT_LLM_KEY or OPENAI_API_KEY
    if not api_key:
        raise HTTPException(500, "AI service not configured")

    import io
    import tempfile
    import subprocess

    # Create meeting record immediately as "processing"
    meeting_record = {
        "id": meeting_id,
        "user_id": user_id,
        "title": meeting_title or f"Meeting {meeting_id[:8]}",
        "participants": participants.split(",") if participants else [],
        "duration_seconds": duration_seconds,
        "status": "transcribing",
        "transcript": None,
        "insights": None,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.meeting_transcripts.update_one(
        {"id": meeting_id},
        {"$set": meeting_record},
        upsert=True,
    )

    contents = await file.read()
    if len(contents) < 1000:
        await db.meeting_transcripts.update_one(
            {"id": meeting_id},
            {"$set": {"status": "failed", "error": "Audio too short or empty"}},
        )
        raise HTTPException(400, "Audio file too short or empty")

    try:

        # If file is large, extract audio with ffmpeg
        audio_data = io.BytesIO(contents)
        audio_data.name = file.filename or "meeting.webm"
        temp_files = []

        if len(contents) > 25 * 1024 * 1024:
            logger.info(f"Large audio ({len(contents)/(1024*1024):.1f}MB), extracting with ffmpeg")
            with tempfile.NamedTemporaryFile(suffix=".webm", delete=False) as tmp_in:
                tmp_in.write(contents)
                temp_files.append(tmp_in.name)
            audio_out = tempfile.mktemp(suffix=".mp3")
            temp_files.append(audio_out)
            result = subprocess.run(
                ["ffmpeg", "-y", "-i", tmp_in.name, "-vn", "-acodec", "libmp3lame",
                 "-ab", "64k", "-ar", "16000", "-ac", "1", audio_out],
                capture_output=True, text=True, timeout=300,
            )
            if result.returncode != 0:
                raise Exception(f"ffmpeg failed: {result.stderr[:200]}")
            audio_data = open(audio_out, "rb")
            audio_data.name = "meeting.mp3"

        # Step 1: Transcribe
        from llm_client import chat_completion, speech_to_text
        logger.info(f"Transcribing meeting {meeting_id}...")
        transcript_response = speech_to_text(audio_file=audio_data, api_key=api_key)
        transcript_text = transcript_response.text
        segments = []
        if hasattr(transcript_response, "segments") and transcript_response.segments:
            segments = [
                {"start": getattr(s, "start", 0), "end": getattr(s, "end", 0), "text": getattr(s, "text", "")}
                for s in transcript_response.segments
            ]

        await db.meeting_transcripts.update_one(
            {"id": meeting_id},
            {"$set": {
                "status": "generating_insights",
                "transcript": {"text": transcript_text, "segments": segments},
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )

        # Step 2: Generate insights with GPT-5.2
        logger.info(f"Generating insights for meeting {meeting_id}...")
        participants_str = ", ".join(meeting_record["participants"]) if meeting_record["participants"] else "Unknown"
        insight_prompt = f"""Analyze this meeting transcript and provide structured insights.

Meeting Title: {meeting_record['title']}
Participants: {participants_str}
Duration: {duration_seconds // 60} minutes

Transcript:
{transcript_text}

Return ONLY valid JSON in this exact format:
{{
  "summary": "2-3 sentence meeting summary",
  "key_decisions": [
    {{"decision": "What was decided", "context": "Brief context"}}
  ],
  "action_items": [
    {{"task": "What needs to be done", "assignee": "Person responsible or 'Unassigned'", "priority": "high|medium|low"}}
  ],
  "topics_discussed": [
    {{"topic": "Topic name", "duration_estimate": "Brief or Detailed", "key_points": ["point1", "point2"]}}
  ],
  "follow_ups": [
    {{"item": "Follow-up item", "due": "Suggested timeline"}}
  ],
  "sentiment": "positive|neutral|mixed|negative",
  "participation_notes": "Brief note on participant engagement"
}}"""

        insight_result = chat_completion(
            messages=[
                {"role": "system", "content": "You are an expert meeting analyst. Return ONLY valid JSON."},
                {"role": "user", "content": insight_prompt},
            ],
            model="gpt-5.5",
            api_key=api_key,
        )
        import json
        raw = insight_result.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
        if raw.endswith("```"):
            raw = raw[:-3]
        if raw.startswith("json"):
            raw = raw[4:]
        insights = json.loads(raw.strip())

        # Save final result
        await db.meeting_transcripts.update_one(
            {"id": meeting_id},
            {"$set": {
                "status": "completed",
                "insights": insights,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )

        # Clean up temp files
        import os as _os
        for f in temp_files:
            try:
                _os.unlink(f)
            except Exception:
                pass

        return {
            "success": True,
            "meeting_id": meeting_id,
            "status": "completed",
            "transcript": {"text": transcript_text, "segments": segments},
            "insights": insights,
        }

    except json.JSONDecodeError:
        await db.meeting_transcripts.update_one(
            {"id": meeting_id},
            {"$set": {"status": "completed", "insights": {"summary": "Insights generation failed. Transcript is available.", "error": True},
                       "updated_at": datetime.now(timezone.utc).isoformat()}},
        )
        return {"success": True, "meeting_id": meeting_id, "status": "completed_partial"}
    except Exception as e:
        logger.error(f"Meeting process error: {e}")
        await db.meeting_transcripts.update_one(
            {"id": meeting_id},
            {"$set": {"status": "failed", "error": str(e), "updated_at": datetime.now(timezone.utc).isoformat()}},
        )
        raise HTTPException(500, f"Processing failed: {str(e)}")


@router.get("/ai/meeting/{meeting_id}/status")
async def get_meeting_status(meeting_id: str):
    """Get the processing status for a meeting transcript."""
    doc = await db.meeting_transcripts.find_one({"id": meeting_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Meeting transcript not found")
    return doc


@router.get("/ai/meeting/user/{user_id}")
async def get_user_meeting_transcripts(user_id: str, limit: int = 50):
    """Get all processed meeting transcripts for a user."""
    docs = await db.meeting_transcripts.find(
        {"user_id": user_id},
        {"_id": 0},
    ).sort("created_at", -1).limit(limit).to_list(limit)
    return {"meetings": docs, "count": len(docs)}


# ============== Meeting Transcript Export ==============

@router.get("/ai/meeting/{meeting_id}/export")
async def export_meeting_transcript(meeting_id: str, format: str = Query("pdf", regex="^(pdf|docx|md)$")):
    """Export meeting transcript + insights as PDF, DOCX, or Markdown."""
    import io

    doc = await db.meeting_transcripts.find_one({"id": meeting_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Meeting transcript not found")
    if doc.get("status") != "completed":
        raise HTTPException(400, "Transcript not ready yet")

    title = doc.get("title", "Meeting Transcript")
    transcript = doc.get("transcript", {}).get("text", "No transcript available")
    insights = doc.get("insights", {})
    created = doc.get("created_at", "")
    participants = doc.get("participants", [])
    duration = doc.get("duration_seconds", 0)
    import re
    safe_name = re.sub(r'[^\w\s-]', '', title)[:60].strip() or "meeting-transcript"

    if format == "md":
        md = _build_meeting_markdown(title, created, participants, duration, transcript, insights)
        return Response(content=md, media_type="text/markdown",
                        headers={"Content-Disposition": f'attachment; filename="{safe_name}.md"'})
    elif format == "pdf":
        pdf_bytes = _build_meeting_pdf(title, created, participants, duration, transcript, insights)
        return Response(content=pdf_bytes, media_type="application/pdf",
                        headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'})
    elif format == "docx":
        docx_bytes = _build_meeting_docx(title, created, participants, duration, transcript, insights)
        return Response(content=docx_bytes,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'})


def _build_meeting_markdown(title, created, participants, duration, transcript, insights):
    lines = [f"# {title}\n"]
    if created:
        lines.append(f"*{created[:10]}*\n")
    if participants:
        lines.append(f"**Participants:** {', '.join(participants)}\n")
    if duration:
        lines.append(f"**Duration:** {duration // 60} minutes\n")
    lines.append("---\n")

    if insights.get("summary"):
        lines.append(f"## Summary\n{insights['summary']}\n")

    if insights.get("key_decisions"):
        lines.append("## Key Decisions\n")
        for d in insights["key_decisions"]:
            lines.append(f"- **{d['decision']}** — {d.get('context', '')}\n")

    if insights.get("action_items"):
        lines.append("## Action Items\n")
        for a in insights["action_items"]:
            lines.append(f"- [ ] {a['task']} *(Assigned to: {a.get('assignee', 'Unassigned')}, Priority: {a.get('priority', 'medium')})*\n")

    if insights.get("follow_ups"):
        lines.append("## Follow-ups\n")
        for f in insights["follow_ups"]:
            lines.append(f"- {f['item']} *(Due: {f.get('due', 'TBD')})*\n")

    lines.append("## Full Transcript\n")
    lines.append(transcript)
    lines.append("\n---\n*Exported from Munal AI*")
    return "\n".join(lines)


def _build_meeting_pdf(title, created, participants, duration, transcript, insights):
    import fitz
    doc = fitz.open()
    W, H = 595, 842
    M = 50
    y = M

    def new_pg():
        nonlocal y
        p = doc.new_page(width=W, height=H)
        y = M
        return p

    def write_text(page, text, fs=10, bold=False, color=(0.15, 0.15, 0.15)):
        nonlocal y
        fn = "helvetica-bold" if bold else "helv"
        for line in text.split('\n'):
            words = line.split(' ')
            current = ""
            for w in words:
                test = f"{current} {w}".strip()
                if len(test) * fs * 0.5 > (W - 2 * M):
                    if y > H - 60:
                        page = new_pg()
                    page.insert_text((M, y), current, fontsize=fs, fontname=fn, color=color)
                    y += fs + 4
                    current = w
                else:
                    current = test
            if current:
                if y > H - 60:
                    page = new_pg()
                page.insert_text((M, y), current, fontsize=fs, fontname=fn, color=color)
                y += fs + 4
        return page

    page = new_pg()
    page.insert_text((M, y), title[:80], fontsize=18, fontname="helvetica-bold", color=(0.29, 0.27, 0.53))
    y += 28
    meta = []
    if created:
        meta.append(created[:10])
    if participants:
        meta.append(f"Participants: {', '.join(participants)}")
    if duration:
        meta.append(f"Duration: {duration // 60} min")
    if meta:
        page.insert_text((M, y), " | ".join(meta), fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))
        y += 16
    page.draw_line((M, y), (W - M, y), color=(0.85, 0.85, 0.85), width=0.5)
    y += 15

    if insights.get("summary"):
        page = write_text(page, "Summary", fs=13, bold=True, color=(0.29, 0.27, 0.53))
        y += 4
        page = write_text(page, insights["summary"])
        y += 10

    if insights.get("action_items"):
        page = write_text(page, "Action Items", fs=13, bold=True, color=(0.29, 0.27, 0.53))
        y += 4
        for a in insights["action_items"]:
            page = write_text(page, f"• {a['task']} [{a.get('assignee', '?')}] ({a.get('priority', 'medium')})")
        y += 10

    if insights.get("key_decisions"):
        page = write_text(page, "Key Decisions", fs=13, bold=True, color=(0.29, 0.27, 0.53))
        y += 4
        for d in insights["key_decisions"]:
            page = write_text(page, f"• {d['decision']}")
        y += 10

    page = write_text(page, "Transcript", fs=13, bold=True, color=(0.29, 0.27, 0.53))
    y += 4
    page = write_text(page, transcript, fs=9, color=(0.3, 0.3, 0.3))

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _build_meeting_docx(title, created, participants, duration, transcript, insights):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(74, 69, 135)

    meta = []
    if created:
        meta.append(created[:10])
    if participants:
        meta.append(f"Participants: {', '.join(participants)}")
    if duration:
        meta.append(f"Duration: {duration // 60} min")
    if meta:
        p = doc.add_paragraph(" | ".join(meta))
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        p.runs[0].font.size = Pt(9)

    if insights.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(insights["summary"])

    if insights.get("action_items"):
        doc.add_heading("Action Items", level=2)
        for a in insights["action_items"]:
            doc.add_paragraph(f"{a['task']} — {a.get('assignee', 'Unassigned')} ({a.get('priority', 'medium')})", style="List Bullet")

    if insights.get("key_decisions"):
        doc.add_heading("Key Decisions", level=2)
        for d in insights["key_decisions"]:
            doc.add_paragraph(d["decision"], style="List Bullet")

    if insights.get("follow_ups"):
        doc.add_heading("Follow-ups", level=2)
        for f in insights["follow_ups"]:
            doc.add_paragraph(f"{f['item']} — Due: {f.get('due', 'TBD')}", style="List Bullet")

    doc.add_heading("Full Transcript", level=2)
    doc.add_paragraph(transcript)
    doc.add_paragraph("Exported from Munal AI").runs[0].font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

