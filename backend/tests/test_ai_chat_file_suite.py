"""
AI Chat File Suite Tests
Tests new file reading + file generation features:
  - _extract_pdf_text, _extract_excel_data, DOCX/text extraction utils
  - _generate_pdf_from_markdown, _generate_docx_from_markdown, _generate_xlsx_from_text
  - Streaming endpoint with file attachments and [GENERATE_*] tags
  - File download endpoint
"""
import os
import sys
import io
import json
import uuid
import time
import pytest
import requests

# Ensure backend on path for direct imports of utilities
sys.path.insert(0, "/app/backend")

def _load_base_url():
    url = os.environ.get("REACT_APP_BACKEND_URL", "").strip()
    if url:
        return url.rstrip("/")
    # Fallback: read /app/frontend/.env
    try:
        with open("/app/frontend/.env", "r") as f:
            for line in f:
                if line.startswith("REACT_APP_BACKEND_URL="):
                    return line.split("=", 1)[1].strip().rstrip("/")
    except Exception:
        pass
    return ""


BASE_URL = _load_base_url()

# ---- Credentials per review request ----
PRIMARY_EMAIL = "chattest@munal.ai"
PRIMARY_PASSWORD = "Test@12345"
FALLBACK_EMAIL = "analytics@munal.ai"
FALLBACK_PASSWORD = "Test@12345"


# =================== Fixtures ===================

@pytest.fixture(scope="session")
def auth_token():
    """Log in (try primary, register if missing, fallback to analytics user)."""
    # Try primary login
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": PRIMARY_EMAIL, "password": PRIMARY_PASSWORD})
    if r.status_code == 200 and "token" in r.json():
        return r.json()["token"]

    # Try register primary
    reg = requests.post(f"{BASE_URL}/api/auth/register", json={
        "email": PRIMARY_EMAIL, "password": PRIMARY_PASSWORD,
        "name": "Chat Test User", "full_name": "Chat Test User",
    })
    if reg.status_code in (200, 201):
        r = requests.post(f"{BASE_URL}/api/auth/login",
                          json={"email": PRIMARY_EMAIL, "password": PRIMARY_PASSWORD})
        if r.status_code == 200 and "token" in r.json():
            return r.json()["token"]

    # Fallback user
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": FALLBACK_EMAIL, "password": FALLBACK_PASSWORD})
    if r.status_code == 200 and "token" in r.json():
        return r.json()["token"]
    pytest.skip(f"Could not authenticate: primary={PRIMARY_EMAIL}, fallback={FALLBACK_EMAIL}")


@pytest.fixture(scope="session")
def headers(auth_token):
    return {"Authorization": f"Bearer {auth_token}", "Content-Type": "application/json"}


@pytest.fixture
def conversation_id(headers):
    r = requests.post(f"{BASE_URL}/api/ai-chat/conversations", headers=headers)
    assert r.status_code == 200, f"Could not create conversation: {r.text}"
    cid = r.json()["id"]
    yield cid
    # Cleanup
    try:
        requests.delete(f"{BASE_URL}/api/ai-chat/conversations/{cid}", headers=headers)
    except Exception:
        pass


# =================== Helper to build test files ===================

def _make_pdf_bytes(text="Hello PDF World. This is a test."):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 720, text)
    c.showPage()
    c.save()
    return buf.getvalue()


def _make_xlsx_bytes():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Age", "City"])
    ws.append(["Alice", 30, "Paris"])
    ws.append(["Bob", 25, "Tokyo"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_docx_bytes():
    from docx import Document
    d = Document()
    d.add_heading("Test Doc", level=1)
    d.add_paragraph("This is a paragraph for DOCX extraction.")
    d.add_paragraph("Second paragraph here.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# =================== Direct Utility Tests ===================

class TestExtractionUtilities:
    """Test the standalone file-extraction utility functions."""

    def test_extract_pdf_text(self):
        from routes.ai_chat import _extract_pdf_text
        pdf_bytes = _make_pdf_bytes("Hello PDF World. This is a test PDF.")
        result = _extract_pdf_text(pdf_bytes)
        assert isinstance(result, str)
        assert "Hello PDF" in result, f"PDF text not extracted: {result[:200]}"
        assert "[Page 1]" in result

    def test_extract_pdf_handles_bad_bytes(self):
        from routes.ai_chat import _extract_pdf_text
        result = _extract_pdf_text(b"not-a-real-pdf")
        # Should not raise, returns an error string
        assert isinstance(result, str)
        assert "Error" in result or "error" in result.lower()

    def test_extract_excel_data(self):
        from routes.ai_chat import _extract_excel_data
        xlsx_bytes = _make_xlsx_bytes()
        result = _extract_excel_data(xlsx_bytes)
        assert isinstance(result, str)
        assert "[Sheet: Sheet1]" in result
        assert "Alice" in result and "Bob" in result
        assert "Paris" in result

    def test_extract_excel_handles_bad_bytes(self):
        from routes.ai_chat import _extract_excel_data
        result = _extract_excel_data(b"not-a-xlsx")
        assert isinstance(result, str)

    def test_extract_docx_via_extract_file_content(self):
        """DOCX extraction lives inside _extract_file_content - test by mocking the storage fetch."""
        import routes.ai_chat as mod
        docx_bytes = _make_docx_bytes()
        # Monkeypatch _get_object so it returns our docx bytes for any path
        original = mod._get_object
        mod._get_object = lambda path: (docx_bytes,
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        try:
            import asyncio
            result_text, img = asyncio.get_event_loop().run_until_complete(
                mod._extract_file_content({
                    "file_id": "fake-id",
                    "filename": "test.docx",
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                })
            )
            assert "[Document: test.docx]" in result_text
            assert "Test Doc" in result_text
            assert "paragraph" in result_text.lower()
            assert img is None
        finally:
            mod._get_object = original

    def test_extract_image_returns_data_url(self):
        """Image attachment should return a base64 data URL for vision API."""
        import routes.ai_chat as mod
        # 1x1 transparent PNG
        png_bytes = (b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
                     b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\rIDATx\x9cc\x00"
                     b"\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82")
        original = mod._get_object
        mod._get_object = lambda path: (png_bytes, "image/png")
        try:
            import asyncio
            text, img_url = asyncio.get_event_loop().run_until_complete(
                mod._extract_file_content({
                    "file_id": "fake-img",
                    "filename": "pixel.png",
                    "content_type": "image/png",
                })
            )
            assert "[Image: pixel.png]" in text
            assert img_url is not None
            assert img_url.startswith("data:image/png;base64,")
        finally:
            mod._get_object = original


class TestGenerationUtilities:
    """Test the file-generation utility functions."""

    def test_generate_pdf_from_markdown(self):
        from routes.ai_chat import _generate_pdf_from_markdown
        md = "# Title\n\n## Subtitle\n\nThis is **bold** and *italic* text.\n\n- Item 1\n- Item 2"
        pdf = _generate_pdf_from_markdown(md)
        assert isinstance(pdf, bytes)
        assert pdf.startswith(b"%PDF"), "Output is not a valid PDF"
        assert len(pdf) > 500

    def test_generate_docx_from_markdown(self):
        from routes.ai_chat import _generate_docx_from_markdown
        from docx import Document
        md = "# Title\n\n## Sub\n\nThis is **bold** content.\n\n- bullet 1\n- bullet 2\n\n1. numbered"
        docx_bytes = _generate_docx_from_markdown(md)
        assert isinstance(docx_bytes, bytes)
        # DOCX = ZIP starts with PK
        assert docx_bytes[:2] == b"PK"
        # Open with python-docx to verify content
        d = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in d.paragraphs)
        assert "Title" in all_text
        assert "bullet 1" in all_text

    def test_generate_xlsx_from_text(self):
        from routes.ai_chat import _generate_xlsx_from_text
        import openpyxl
        md = "| Name | Age |\n| --- | --- |\n| Alice | 30 |\n| Bob | 25 |"
        xlsx_bytes = _generate_xlsx_from_text(md)
        assert isinstance(xlsx_bytes, bytes)
        assert xlsx_bytes[:2] == b"PK"
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb.active
        # Collect all cell values
        values = []
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if v is not None:
                    values.append(str(v))
        assert "Name" in values
        assert "Alice" in values
        assert "30" in values


# =================== Streaming endpoint integration tests ===================

class TestStreamingWithAttachments:
    """Send messages with attachments / generation tags through the live streaming endpoint."""

    def _consume_sse(self, response, timeout=120):
        """Collect all SSE events from a streaming response."""
        events = []
        start = time.time()
        for line in response.iter_lines(decode_unicode=True):
            if time.time() - start > timeout:
                break
            if not line:
                continue
            if line.startswith("data: "):
                try:
                    events.append(json.loads(line[6:]))
                except json.JSONDecodeError:
                    pass
        return events

    def _upload(self, headers, filename, content_type, data):
        files = {"file": (filename, data, content_type)}
        # /upload uses multipart - strip Content-Type from auth headers
        h = {k: v for k, v in headers.items() if k.lower() != "content-type"}
        r = requests.post(f"{BASE_URL}/api/ai-chat/upload", headers=h, files=files, timeout=60)
        assert r.status_code == 200, f"Upload failed: {r.status_code} {r.text}"
        return r.json()

    def test_send_message_with_pdf_attachment(self, headers, conversation_id):
        """Upload a PDF, attach it to a message, verify streaming completes."""
        pdf = _make_pdf_bytes("The capital of France is Paris.")
        rec = self._upload(headers, "test.pdf", "application/pdf", pdf)
        attachment = {
            "file_id": rec["storage_path"].split("/")[-1].rsplit(".", 1)[0],
            "filename": rec["original_filename"],
            "content_type": rec["content_type"],
            "storage_path": rec["storage_path"],
        }
        # Note: /upload stores under storage_path; _extract_file_content reads from
        # ai-chat-files/{file_id}. The attachment shape may differ - use what frontend sends.
        # Use the full storage_path-based file_id (uuid before extension).
        body = {
            "content": "What does the attached document say?",
            "attachments": [attachment],
        }
        r = requests.post(f"{BASE_URL}/api/ai-chat/conversations/{conversation_id}/messages",
                          headers=headers, json=body, stream=True, timeout=120)
        assert r.status_code == 200, f"Send failed: {r.status_code} {r.text}"
        events = self._consume_sse(r)
        types = [e.get("type") for e in events]
        assert "thinking" in types, f"missing thinking, got {types[:10]}"
        assert "done" in types, f"missing done, got {types[:10]}"
        # Ensure chunks present
        chunks = [e for e in events if e.get("type") == "chunk"]
        assert len(chunks) > 0, "no chunk events received"

    def test_send_message_triggers_pdf_generation(self, headers, conversation_id):
        """Ask AI for a PDF - expect [GENERATE_PDF] tag and generated_files in done event."""
        body = {"content": "Please create a PDF document with a heading 'Munal Report' and one paragraph "
                           "about productivity. End your message with [GENERATE_PDF] so the system "
                           "converts it to a downloadable PDF."}
        r = requests.post(f"{BASE_URL}/api/ai-chat/conversations/{conversation_id}/messages",
                          headers=headers, json=body, stream=True, timeout=180)
        assert r.status_code == 200
        events = self._consume_sse(r, timeout=180)
        done = next((e for e in events if e.get("type") == "done"), None)
        assert done is not None, "No done event received"
        gen = done.get("generated_files", [])
        full_text = "".join(e.get("content", "") for e in events if e.get("type") == "chunk")
        # Either gen contains a pdf entry, or the AI didn't include the tag (acceptable)
        if "[GENERATE_PDF]" in full_text:
            assert any(f.get("type") == "pdf" for f in gen), \
                f"Tag present but no pdf in generated_files: {gen}"
        else:
            pytest.skip(f"AI response did not include [GENERATE_PDF] tag; skipping. Text: {full_text[:200]}")

    def test_send_message_triggers_xlsx_generation(self, headers, conversation_id):
        body = {"content": "Create a small spreadsheet showing this data:\n"
                           "| Name | Score |\n| --- | --- |\n| Alice | 90 |\n| Bob | 85 |\n"
                           "End your reply with [GENERATE_XLSX] so we get an Excel file."}
        r = requests.post(f"{BASE_URL}/api/ai-chat/conversations/{conversation_id}/messages",
                          headers=headers, json=body, stream=True, timeout=180)
        assert r.status_code == 200
        events = self._consume_sse(r, timeout=180)
        done = next((e for e in events if e.get("type") == "done"), None)
        assert done is not None
        full_text = "".join(e.get("content", "") for e in events if e.get("type") == "chunk")
        gen = done.get("generated_files", [])
        if "[GENERATE_XLSX]" in full_text:
            assert any(f.get("type") == "xlsx" for f in gen), \
                f"XLSX tag present but no xlsx in generated_files: {gen}"
            # Confirm download URL pattern
            xlsx_entry = next(f for f in gen if f.get("type") == "xlsx")
            assert xlsx_entry.get("url", "").startswith("/api/ai-chat/files/")
        else:
            pytest.skip("AI response did not include [GENERATE_XLSX] tag")

    def test_send_message_triggers_docx_generation(self, headers, conversation_id):
        body = {"content": "Write a short Word document with title 'Hello' and a paragraph. "
                           "End your reply with [GENERATE_DOCX]."}
        r = requests.post(f"{BASE_URL}/api/ai-chat/conversations/{conversation_id}/messages",
                          headers=headers, json=body, stream=True, timeout=180)
        assert r.status_code == 200
        events = self._consume_sse(r, timeout=180)
        done = next((e for e in events if e.get("type") == "done"), None)
        assert done is not None
        full_text = "".join(e.get("content", "") for e in events if e.get("type") == "chunk")
        gen = done.get("generated_files", [])
        if "[GENERATE_DOCX]" in full_text:
            assert any(f.get("type") == "docx" for f in gen), \
                f"DOCX tag present but no docx in generated_files: {gen}"
        else:
            pytest.skip("AI response did not include [GENERATE_DOCX] tag")


# =================== Download endpoint tests ===================

class TestFileDownload:
    """Verify the file-download endpoint."""

    def test_download_nonexistent_file_returns_404(self, headers):
        """Unknown file id should return 404."""
        r = requests.get(f"{BASE_URL}/api/ai-chat/files/{uuid.uuid4()}", headers=headers)
        assert r.status_code == 404, f"Expected 404, got {r.status_code}: {r.text[:200]}"

    def test_download_route_conflict_check(self):
        """RCA check: ai_chat.py defines two GET /files/{file_id} routes.

        Line ~855 is auth-protected and queries db.ai_chat_files.
        Line ~1032 (download_generated_file) has no auth and reads object storage directly.
        FastAPI uses the FIRST matching route -> the second (generated-file) route is
        unreachable. Generated PDF/DOCX/XLSX/PNG cannot be downloaded by their
        returned URL because the first handler looks them up in ai_chat_files collection.
        """
        from routes import ai_chat as mod
        files_routes = []
        for r in mod.router.routes:
            path = getattr(r, "path", "")
            methods = getattr(r, "methods", set()) or set()
            if path.endswith("/files/{file_id}") and "GET" in methods:
                files_routes.append(r)
        print(f"Found {len(files_routes)} GET /files/{{file_id}} routes: "
              f"{[r.endpoint.__name__ for r in files_routes]}")
        assert len(files_routes) >= 2, (
            f"Expected at least 2 duplicate /files/{{file_id}} GET routes, "
            f"found {len(files_routes)}"
        )


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
