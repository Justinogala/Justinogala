"""
Test conversion history and Word-to-PDF/PDF-to-Word endpoints with mammoth+weasyprint
Tests for iteration 40: Fixed Word-to-PDF converter and Conversion History feature
"""
import pytest
import requests
import os
import io
import uuid

# Create test DOCX file using python-docx
from docx import Document

# Create test PDF using reportlab
from reportlab.pdfgen import canvas

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

# Generate unique user_id for test isolation
TEST_USER_ID = f"test-conversion-{uuid.uuid4().hex[:8]}"


@pytest.fixture(scope="module")
def test_docx_file():
    """Create a test DOCX file in memory."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph for conversion testing.')
    doc.add_paragraph('Second paragraph with more content for testing mammoth+weasyprint.')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture(scope="module")
def test_pdf_file():
    """Create a test PDF file in memory."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 730, "This is a test PDF for conversion.")
    c.drawString(100, 710, "Testing PDF to Word conversion.")
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


class TestWordToPdfConversion:
    """Tests for Word to PDF conversion using mammoth+weasyprint"""

    def test_convert_docx_to_pdf_success(self, test_docx_file):
        """POST /api/esignature/convert-to-pdf converts DOCX to valid PDF"""
        files = {'file': ('test_document.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'user_id': TEST_USER_ID}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files, data=data)
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
        assert response.headers.get('Content-Type') == 'application/pdf'
        
        # Verify PDF content starts with %PDF- header
        pdf_content = response.content
        assert pdf_content.startswith(b'%PDF-'), "Response should be a valid PDF (starts with %PDF-)"
        assert len(pdf_content) > 100, "PDF should have substantial content"
        print(f"✓ Word to PDF conversion successful - PDF size: {len(pdf_content)} bytes")

    def test_convert_docx_without_user_id(self, test_docx_file):
        """POST /api/esignature/convert-to-pdf works without user_id (no history saved)"""
        files = {'file': ('test_no_history.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files)
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        assert response.content.startswith(b'%PDF-'), "Response should be a valid PDF"
        print("✓ Word to PDF conversion works without user_id")

    def test_reject_pdf_file(self, test_pdf_file):
        """POST /api/esignature/convert-to-pdf rejects PDF files with 400"""
        files = {'file': ('test.pdf', test_pdf_file, 'application/pdf')}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files)
        
        assert response.status_code == 400, f"Expected 400, got {response.status_code}"
        data = response.json()
        assert 'detail' in data
        assert 'DOC' in data['detail'] or 'DOCX' in data['detail'] or 'supported' in data['detail'].lower()
        print("✓ PDF files correctly rejected with 400 error")

    def test_reject_txt_file(self):
        """POST /api/esignature/convert-to-pdf rejects TXT files with 400"""
        files = {'file': ('test.txt', b'This is plain text', 'text/plain')}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files)
        
        assert response.status_code == 400, f"Expected 400, got {response.status_code}"
        print("✓ TXT files correctly rejected with 400 error")


class TestPdfToWordConversion:
    """Tests for PDF to Word conversion"""

    def test_convert_pdf_to_word_success(self, test_pdf_file):
        """POST /api/esignature/convert-to-word converts PDF to DOCX"""
        files = {'file': ('test_document.pdf', test_pdf_file, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-word", files=files, data=data)
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
        content_type = response.headers.get('Content-Type', '')
        assert 'officedocument' in content_type or 'application/octet' in content_type
        
        # DOCX files start with PK (ZIP format)
        assert response.content[:2] == b'PK', "Response should be a valid DOCX (ZIP format)"
        print(f"✓ PDF to Word conversion successful - DOCX size: {len(response.content)} bytes")

    def test_convert_pdf_without_user_id(self, test_pdf_file):
        """POST /api/esignature/convert-to-word works without user_id"""
        files = {'file': ('test_no_history.pdf', test_pdf_file, 'application/pdf')}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-word", files=files)
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        assert response.content[:2] == b'PK', "Response should be a valid DOCX"
        print("✓ PDF to Word conversion works without user_id")

    def test_reject_docx_file(self, test_docx_file):
        """POST /api/esignature/convert-to-word rejects DOCX files with 400"""
        files = {'file': ('test.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        
        response = requests.post(f"{BASE_URL}/api/esignature/convert-to-word", files=files)
        
        assert response.status_code == 400, f"Expected 400, got {response.status_code}"
        data = response.json()
        assert 'detail' in data
        assert 'PDF' in data['detail'] or 'supported' in data['detail'].lower()
        print("✓ DOCX files correctly rejected with 400 error")


class TestConversionHistory:
    """Tests for conversion history endpoints"""

    def test_get_conversion_history_empty(self):
        """GET /api/esignature/conversion-history returns empty list for new user"""
        new_user_id = f"test-new-user-{uuid.uuid4().hex[:8]}"
        
        response = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': new_user_id})
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert 'history' in data
        assert isinstance(data['history'], list)
        print("✓ Empty conversion history returns correctly")

    def test_get_conversion_history_with_entries(self, test_docx_file, test_pdf_file):
        """GET /api/esignature/conversion-history returns list of conversions (without file_data)"""
        user_id = f"test-history-{uuid.uuid4().hex[:8]}"
        
        # First, create a conversion to have history
        files = {'file': ('history_test.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'user_id': user_id}
        convert_resp = requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files, data=data)
        assert convert_resp.status_code == 200, "Conversion should succeed"
        
        # Now get history
        response = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': user_id})
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert 'history' in data
        assert len(data['history']) >= 1, "Should have at least one history entry"
        
        entry = data['history'][0]
        assert 'id' in entry
        assert 'user_id' in entry
        assert 'conversion_type' in entry
        assert 'original_filename' in entry
        assert 'converted_filename' in entry
        assert 'original_size' in entry
        assert 'converted_size' in entry
        assert 'created_at' in entry
        assert 'file_data' not in entry, "file_data should be excluded for performance"
        
        assert entry['conversion_type'] == 'word-to-pdf'
        assert entry['original_filename'] == 'history_test.docx'
        assert entry['converted_filename'] == 'history_test.pdf'
        print(f"✓ Conversion history contains correct entry: {entry['original_filename']} -> {entry['converted_filename']}")
        
        # Store for cleanup
        return entry['id'], user_id

    def test_download_conversion(self, test_docx_file):
        """GET /api/esignature/conversion-history/{id}/download returns the converted file"""
        user_id = f"test-download-{uuid.uuid4().hex[:8]}"
        
        # Create a conversion
        files = {'file': ('download_test.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'user_id': user_id}
        requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files, data=data)
        
        # Get the history entry ID
        history_resp = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': user_id})
        entry_id = history_resp.json()['history'][0]['id']
        
        # Download
        response = requests.get(f"{BASE_URL}/api/esignature/conversion-history/{entry_id}/download")
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        assert response.headers.get('Content-Type') == 'application/pdf'
        assert response.content.startswith(b'%PDF-'), "Downloaded file should be valid PDF"
        print(f"✓ Conversion download works - file size: {len(response.content)} bytes")

    def test_download_conversion_not_found(self):
        """GET /api/esignature/conversion-history/{id}/download returns 404 for invalid ID"""
        fake_id = f"nonexistent-{uuid.uuid4().hex}"
        
        response = requests.get(f"{BASE_URL}/api/esignature/conversion-history/{fake_id}/download")
        
        assert response.status_code == 404, f"Expected 404, got {response.status_code}"
        print("✓ Download returns 404 for nonexistent entry")

    def test_delete_conversion(self, test_docx_file):
        """DELETE /api/esignature/conversion-history/{id} removes entry"""
        user_id = f"test-delete-{uuid.uuid4().hex[:8]}"
        
        # Create a conversion
        files = {'file': ('delete_test.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'user_id': user_id}
        requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files, data=data)
        
        # Get the history entry ID
        history_resp = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': user_id})
        entry_id = history_resp.json()['history'][0]['id']
        
        # Delete
        response = requests.delete(f"{BASE_URL}/api/esignature/conversion-history/{entry_id}")
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert data.get('success') == True
        
        # Verify deletion
        history_resp2 = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': user_id})
        assert len(history_resp2.json()['history']) == 0, "History should be empty after delete"
        print("✓ Conversion delete works correctly")

    def test_delete_conversion_not_found(self):
        """DELETE /api/esignature/conversion-history/{id} returns 404 for invalid ID"""
        fake_id = f"nonexistent-{uuid.uuid4().hex}"
        
        response = requests.delete(f"{BASE_URL}/api/esignature/conversion-history/{fake_id}")
        
        assert response.status_code == 404, f"Expected 404, got {response.status_code}"
        print("✓ Delete returns 404 for nonexistent entry")


class TestConversionHistorySavesBoth:
    """Tests that both Word-to-PDF and PDF-to-Word save to history"""

    def test_word_to_pdf_saves_history(self, test_docx_file):
        """Word to PDF conversion saves to conversion_history with user_id"""
        user_id = f"test-w2p-history-{uuid.uuid4().hex[:8]}"
        
        files = {'file': ('w2p_history.docx', test_docx_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'user_id': user_id}
        
        # Convert
        requests.post(f"{BASE_URL}/api/esignature/convert-to-pdf", files=files, data=data)
        
        # Check history
        history_resp = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': user_id})
        history = history_resp.json()['history']
        
        assert len(history) == 1, "Should have 1 history entry"
        assert history[0]['conversion_type'] == 'word-to-pdf'
        print("✓ Word to PDF saves to conversion history")

    def test_pdf_to_word_saves_history(self, test_pdf_file):
        """PDF to Word conversion saves to conversion_history with user_id"""
        user_id = f"test-p2w-history-{uuid.uuid4().hex[:8]}"
        
        files = {'file': ('p2w_history.pdf', test_pdf_file, 'application/pdf')}
        data = {'user_id': user_id}
        
        # Convert
        requests.post(f"{BASE_URL}/api/esignature/convert-to-word", files=files, data=data)
        
        # Check history
        history_resp = requests.get(f"{BASE_URL}/api/esignature/conversion-history", params={'user_id': user_id})
        history = history_resp.json()['history']
        
        assert len(history) == 1, "Should have 1 history entry"
        assert history[0]['conversion_type'] == 'pdf-to-word'
        print("✓ PDF to Word saves to conversion history")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
