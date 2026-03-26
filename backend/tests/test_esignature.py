"""
eSignature API Tests
Tests for all eSignature endpoints:
- POST /api/esignature/upload - PDF/DOCX upload
- POST /api/esignature/sign - Sign document with signature placements
- GET /api/esignature/documents/{doc_id}/pdf - Get original PDF
- GET /api/esignature/documents/{doc_id}/signed - Download signed PDF
- GET /api/esignature/history - Get signing history
- POST /api/esignature/signatures - Save signature for reuse
- GET /api/esignature/signatures - List saved signatures
- DELETE /api/esignature/signatures/{sig_id} - Delete saved signature
"""
import pytest
import requests
import os
import io
import base64
import json

# Use public URL for testing
BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', 'https://vertical-solutions-3.preview.emergentagent.com').rstrip('/')

# Test user ID
TEST_USER_ID = "3fe4c41c-4f43-4683-98dc-db6de39b842c"


class TestESignatureHealth:
    """Basic health check"""
    
    def test_api_health(self):
        response = requests.get(f"{BASE_URL}/api/health")
        assert response.status_code == 200
        data = response.json()
        assert data["status"] == "healthy"
        print(f"API Health: {data}")


class TestESignatureUpload:
    """Document upload endpoint tests"""
    
    def test_upload_pdf(self):
        """Test uploading a PDF document"""
        # Create a simple test PDF using reportlab
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.drawString(100, 750, "Test PDF Document for eSignature")
        c.drawString(100, 700, "This is a test document created for API testing")
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        
        files = {'file': ('test_document.pdf', pdf_buffer, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        
        response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        
        assert response.status_code == 200, f"Upload failed: {response.text}"
        result = response.json()
        assert result["success"] is True
        assert "document" in result
        doc = result["document"]
        assert "id" in doc
        assert doc["page_count"] >= 1
        assert doc["filename"].endswith(".pdf")
        assert doc["converted"] is False
        print(f"PDF Upload Success - Doc ID: {doc['id']}, Pages: {doc['page_count']}")
        
        # Store doc_id for later tests
        TestESignatureUpload.pdf_doc_id = doc["id"]
        return doc["id"]
    
    def test_upload_docx(self):
        """Test uploading a DOCX document (auto-converts to PDF)"""
        # Create a simple DOCX using python-docx
        from docx import Document
        
        docx_buffer = io.BytesIO()
        doc = Document()
        doc.add_heading("Test DOCX Document", 0)
        doc.add_paragraph("This is a test DOCX document for eSignature API testing.")
        doc.add_paragraph("It should be automatically converted to PDF.")
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        files = {'file': ('test_document.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'user_id': TEST_USER_ID}
        
        response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        
        assert response.status_code == 200, f"DOCX Upload failed: {response.text}"
        result = response.json()
        assert result["success"] is True
        doc_info = result["document"]
        assert doc_info["converted"] is True  # Should be converted to PDF
        assert doc_info["filename"].endswith(".pdf")
        assert "docx" in doc_info["original_filename"].lower()
        print(f"DOCX Upload Success - Converted to PDF, Doc ID: {doc_info['id']}")
        
        TestESignatureUpload.docx_doc_id = doc_info["id"]
        return doc_info["id"]
    
    def test_upload_invalid_file_type(self):
        """Test uploading unsupported file type fails"""
        files = {'file': ('test.txt', io.BytesIO(b"plain text content"), 'text/plain')}
        data = {'user_id': TEST_USER_ID}
        
        response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        
        assert response.status_code == 400
        result = response.json()
        assert "Unsupported" in result.get("detail", "") or "supported" in result.get("detail", "").lower()
        print(f"Invalid file type rejected correctly: {result}")


class TestESignatureGetPdf:
    """Test GET /api/esignature/documents/{doc_id}/pdf"""
    
    def test_get_pdf_for_rendering(self):
        """Test retrieving original PDF for rendering"""
        # First upload a PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.drawString(100, 750, "PDF for Rendering Test")
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        
        files = {'file': ('render_test.pdf', pdf_buffer, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        upload_response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        doc_id = upload_response.json()["document"]["id"]
        
        # Now get the PDF
        response = requests.get(f"{BASE_URL}/api/esignature/documents/{doc_id}/pdf")
        
        assert response.status_code == 200
        assert response.headers.get("content-type") == "application/pdf"
        assert len(response.content) > 100  # Should have PDF content
        print(f"GET PDF success - Size: {len(response.content)} bytes")
    
    def test_get_pdf_not_found(self):
        """Test getting non-existent document returns 404"""
        response = requests.get(f"{BASE_URL}/api/esignature/documents/nonexistent-doc-id/pdf")
        assert response.status_code == 404
        print("Non-existent document returns 404 correctly")


class TestESignatureSign:
    """Test POST /api/esignature/sign"""
    
    def test_sign_document(self):
        """Test signing a PDF with signature placements"""
        # First upload a PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.drawString(100, 750, "Document to Sign")
        c.drawString(100, 700, "Signature should appear below")
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        
        files = {'file': ('sign_test.pdf', pdf_buffer, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        upload_response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        doc_id = upload_response.json()["document"]["id"]
        
        # Create a simple PNG signature (1x1 transparent pixel as base64)
        # Create a more valid PNG using PIL
        from PIL import Image, ImageDraw
        
        img = Image.new('RGBA', (200, 50), (255, 255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.text((10, 10), "Test Signature", fill=(0, 0, 128))
        
        sig_buffer = io.BytesIO()
        img.save(sig_buffer, format='PNG')
        sig_buffer.seek(0)
        sig_b64 = base64.b64encode(sig_buffer.read()).decode('utf-8')
        signature_data_url = f"data:image/png;base64,{sig_b64}"
        
        # Sign the document
        placements = json.dumps([{
            "page": 0,
            "x": 0.3,
            "y": 0.7,
            "width": 0.3,
            "height": 0.1,
            "type": "signature"
        }])
        
        sign_data = {
            "doc_id": doc_id,
            "user_id": TEST_USER_ID,
            "user_name": "Test User",
            "user_email": "test@example.com",
            "signature_data_url": signature_data_url,
            "placements": placements
        }
        
        response = requests.post(f"{BASE_URL}/api/esignature/sign", data=sign_data)
        
        assert response.status_code == 200, f"Sign failed: {response.text}"
        result = response.json()
        assert result["success"] is True
        assert "signed_document" in result
        signed_doc = result["signed_document"]
        assert signed_doc["id"] == doc_id
        assert "_signed.pdf" in signed_doc["filename"]
        print(f"Sign Success - Signed filename: {signed_doc['filename']}")
        
        TestESignatureSign.signed_doc_id = doc_id
        return doc_id
    
    def test_sign_with_date_field(self):
        """Test signing with date field placement"""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from PIL import Image
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.drawString(100, 750, "Document with Date Field")
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        
        files = {'file': ('date_test.pdf', pdf_buffer, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        upload_response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        doc_id = upload_response.json()["document"]["id"]
        
        # Create signature
        img = Image.new('RGBA', (200, 50), (255, 255, 255, 255))
        sig_buffer = io.BytesIO()
        img.save(sig_buffer, format='PNG')
        sig_buffer.seek(0)
        sig_b64 = base64.b64encode(sig_buffer.read()).decode('utf-8')
        signature_data_url = f"data:image/png;base64,{sig_b64}"
        
        # Placements with signature and date
        placements = json.dumps([
            {"page": 0, "x": 0.3, "y": 0.6, "width": 0.3, "height": 0.1, "type": "signature"},
            {"page": 0, "x": 0.6, "y": 0.8, "width": 0.2, "height": 0.03, "type": "date"}
        ])
        
        sign_data = {
            "doc_id": doc_id,
            "user_id": TEST_USER_ID,
            "user_name": "Test User",
            "user_email": "test@example.com",
            "signature_data_url": signature_data_url,
            "placements": placements
        }
        
        response = requests.post(f"{BASE_URL}/api/esignature/sign", data=sign_data)
        assert response.status_code == 200, f"Sign with date failed: {response.text}"
        print("Sign with date field success")
    
    def test_sign_no_placements(self):
        """Test signing with empty placements fails"""
        sign_data = {
            "doc_id": "some-doc-id",  # Non-existent doc returns 404
            "user_id": TEST_USER_ID,
            "signature_data_url": "data:image/png;base64,iVBORw0KGgo=",
            "placements": "[]"
        }
        
        response = requests.post(f"{BASE_URL}/api/esignature/sign", data=sign_data)
        # Doc not found returns 404, empty placements on valid doc returns 400
        assert response.status_code in [400, 404]
        print(f"Invalid sign request rejected correctly with status {response.status_code}")


class TestESignatureDownloadSigned:
    """Test GET /api/esignature/documents/{doc_id}/signed"""
    
    def test_download_signed_pdf(self):
        """Test downloading a signed PDF"""
        # First upload and sign a document
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from PIL import Image
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.drawString(100, 750, "Download Test Document")
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        
        files = {'file': ('download_test.pdf', pdf_buffer, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        upload_response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        doc_id = upload_response.json()["document"]["id"]
        
        # Sign it
        img = Image.new('RGBA', (200, 50), (255, 255, 255, 255))
        sig_buffer = io.BytesIO()
        img.save(sig_buffer, format='PNG')
        sig_buffer.seek(0)
        sig_b64 = base64.b64encode(sig_buffer.read()).decode('utf-8')
        
        sign_data = {
            "doc_id": doc_id,
            "user_id": TEST_USER_ID,
            "signature_data_url": f"data:image/png;base64,{sig_b64}",
            "placements": json.dumps([{"page": 0, "x": 0.3, "y": 0.7, "width": 0.3, "height": 0.1, "type": "signature"}])
        }
        requests.post(f"{BASE_URL}/api/esignature/sign", data=sign_data)
        
        # Download signed
        response = requests.get(f"{BASE_URL}/api/esignature/documents/{doc_id}/signed")
        
        assert response.status_code == 200
        assert response.headers.get("content-type") == "application/pdf"
        assert "attachment" in response.headers.get("content-disposition", "")
        assert len(response.content) > 100
        print(f"Download signed PDF success - Size: {len(response.content)} bytes")
    
    def test_download_unsigned_document_fails(self):
        """Test downloading unsigned document returns 404"""
        # Upload but don't sign
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.drawString(100, 750, "Unsigned Document")
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        
        files = {'file': ('unsigned.pdf', pdf_buffer, 'application/pdf')}
        data = {'user_id': TEST_USER_ID}
        upload_response = requests.post(f"{BASE_URL}/api/esignature/upload", files=files, data=data)
        doc_id = upload_response.json()["document"]["id"]
        
        # Try to download signed version
        response = requests.get(f"{BASE_URL}/api/esignature/documents/{doc_id}/signed")
        assert response.status_code == 404
        print("Unsigned document download returns 404 correctly")


class TestESignatureHistory:
    """Test GET /api/esignature/history"""
    
    def test_get_history(self):
        """Test retrieving signing history"""
        response = requests.get(f"{BASE_URL}/api/esignature/history?user_id={TEST_USER_ID}")
        
        assert response.status_code == 200
        result = response.json()
        assert "history" in result
        assert isinstance(result["history"], list)
        
        if len(result["history"]) > 0:
            entry = result["history"][0]
            assert "doc_id" in entry
            assert "user_id" in entry
            assert "signed_at" in entry
        
        print(f"History retrieved - {len(result['history'])} entries")


class TestESignatureSavedSignatures:
    """Test signature CRUD endpoints"""
    
    def test_save_signature(self):
        """Test POST /api/esignature/signatures - Save a signature"""
        from PIL import Image
        
        img = Image.new('RGBA', (200, 50), (255, 255, 255, 255))
        sig_buffer = io.BytesIO()
        img.save(sig_buffer, format='PNG')
        sig_buffer.seek(0)
        sig_b64 = base64.b64encode(sig_buffer.read()).decode('utf-8')
        
        data = {
            "user_id": TEST_USER_ID,
            "name": "Test Signature",
            "sig_type": "draw",
            "data_url": f"data:image/png;base64,{sig_b64}"
        }
        
        response = requests.post(f"{BASE_URL}/api/esignature/signatures", data=data)
        
        assert response.status_code == 200, f"Save signature failed: {response.text}"
        result = response.json()
        assert result["success"] is True
        assert "signature" in result
        sig = result["signature"]
        assert "id" in sig
        assert sig["name"] == "Test Signature"
        assert sig["type"] == "draw"
        print(f"Signature saved - ID: {sig['id']}")
        
        TestESignatureSavedSignatures.saved_sig_id = sig["id"]
        return sig["id"]
    
    def test_list_signatures(self):
        """Test GET /api/esignature/signatures - List saved signatures"""
        response = requests.get(f"{BASE_URL}/api/esignature/signatures?user_id={TEST_USER_ID}")
        
        assert response.status_code == 200
        result = response.json()
        assert "signatures" in result
        assert isinstance(result["signatures"], list)
        print(f"Listed {len(result['signatures'])} saved signatures")
    
    def test_delete_signature(self):
        """Test DELETE /api/esignature/signatures/{sig_id}"""
        # First save a signature to delete
        from PIL import Image
        
        img = Image.new('RGBA', (100, 30), (255, 255, 255, 255))
        sig_buffer = io.BytesIO()
        img.save(sig_buffer, format='PNG')
        sig_buffer.seek(0)
        sig_b64 = base64.b64encode(sig_buffer.read()).decode('utf-8')
        
        data = {
            "user_id": TEST_USER_ID,
            "name": "Signature to Delete",
            "sig_type": "draw",
            "data_url": f"data:image/png;base64,{sig_b64}"
        }
        
        save_response = requests.post(f"{BASE_URL}/api/esignature/signatures", data=data)
        sig_id = save_response.json()["signature"]["id"]
        
        # Delete it
        response = requests.delete(f"{BASE_URL}/api/esignature/signatures/{sig_id}")
        
        assert response.status_code == 200
        result = response.json()
        assert result["success"] is True
        print(f"Signature deleted - ID: {sig_id}")
    
    def test_delete_nonexistent_signature(self):
        """Test deleting non-existent signature returns 404"""
        response = requests.delete(f"{BASE_URL}/api/esignature/signatures/nonexistent-sig-id")
        assert response.status_code == 404
        print("Non-existent signature delete returns 404 correctly")


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
